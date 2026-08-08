#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
=====================================================================================
 MATRIX SCANNER - Tableau de bord pédagogique de reconnaissance réseau et de CVE
 Concepteur & Développeur : Ronan BOZOC
=====================================================================================

  ⚠️  AVERTISSEMENT LÉGAL ET ÉTHIQUE  ⚠️
  -----------------------------------------------------------------------------------
  Cet outil exécute des scans réseau réels (nmap) sur le système hôte.
  Le scan d'un système informatique sans autorisation explicite et écrite de son
  propriétaire est ILLÉGAL dans la quasi-totalité des juridictions (en France :
  article 323-1 et suivants du Code pénal - accès ou maintien frauduleux dans un
  système de traitement automatisé de données).

  Cette application est fournie à des fins strictement PÉDAGOGIQUES, pour être
  utilisée UNIQUEMENT sur :
    - des machines virtuelles personnelles (ex: Metasploitable, VulnHub, TryHackMe),
    - des réseaux de laboratoire isolés,
    - des cibles pour lesquelles vous disposez d'une autorisation écrite (pentest).

  Ne scannez JAMAIS une cible dont vous n'êtes pas propriétaire ou pour laquelle
  vous n'avez pas reçu d'autorisation explicite.
=====================================================================================

Architecture :
    - Backend Flask exposant une API REST (/api/...)
    - Communication avec l'OS via le module `subprocess` (jamais de shell=True avec
      une entrée utilisateur non validée -> protection contre l'injection de commande)
    - Parsing des résultats nmap au format XML (sortie native de nmap, -oX -)
    - Élévation de privilèges pour l'installation de paquets via `pkexec`, qui
      affiche une pop-up graphique Polkit native (aucun mot de passe ne transite
      par le code Python : c'est Polkit qui gère l'authentification système).
"""

import base64
import io
import importlib.util
import ipaddress
import json
import psutil
import re
import shutil
import subprocess
import threading
import time
import uuid
import tempfile
import os
import xml.etree.ElementTree as ET
from datetime import datetime
import urllib.request
from urllib.parse import urlparse
import socket
import concurrent.futures
import sqlite3
from flask import Flask, jsonify, render_template, request, send_file

DB_PATH = os.path.join(os.path.dirname(__file__), "scanner_history.db")

def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS history (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                tool TEXT,
                target TEXT,
                mode TEXT DEFAULT 'distant',
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                result_json TEXT
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS custom_reports (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                author TEXT DEFAULT 'Auditeur Cybersécurité',
                description TEXT DEFAULT '',
                content_html TEXT DEFAULT '',
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS report_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                report_id INTEGER,
                item_title TEXT,
                item_html TEXT,
                added_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(report_id) REFERENCES custom_reports(id) ON DELETE CASCADE
            )
        ''')
        conn.commit()

def save_history(tool, target, result, mode="distant"):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("INSERT INTO history (tool, target, mode, result_json) VALUES (?, ?, ?, ?)",
                         (tool, target, mode, json.dumps(result)))
            conn.commit()
    except Exception as e:
        print(f"Erreur sauvegarde SQLite : {e}")

init_db()

app = Flask(__name__)

# ---------------------------------------------------------------------------------
# État global très simple (application mono-utilisateur, usage local uniquement)
# ---------------------------------------------------------------------------------
INSTALL_STATUS = {
    "running": False,
    "success": None,
    "log": "",
    "finished_at": None,
}

SCAN_JOBS = {}  # job_id -> dict(status, result, error)
NUCLEI_JOBS = {}  # job_id -> dict(status, result, error)
WHATWEB_JOBS = {}  # job_id -> dict(status, result, error)
ACTIVE_PROCESSES = {}  # job_id -> subprocess.Popen

# Paquets système nécessaires et leur binaire/commande de vérification associée
REQUIRED_TOOLS = {
    "nmap": {
        "check_cmd": ["nmap", "--version"],
        "apt_package": "nmap",
        "category": "Scans Réseau & Exploration LAN",
        "description": "Scanner de ports et de services réseau",
    },
    "searchsploit": {
        "check_cmd": ["searchsploit", "-h"],
        "apt_package": "exploitdb",
        "category": "Fuzzing & Recherche de Vulnérabilités",
        "description": "Recherche d'exploits/CVE connus (base Exploit-DB locale)",
    },
    "nuclei": {
        "check_cmd": ["nuclei", "-version"],
        "apt_package": "nuclei",
        "category": "Fuzzing & Recherche de Vulnérabilités",
        "description": "Scanner de vulnérabilités basé sur des templates (web/réseau)",
    },
    "whatweb": {
        "check_cmd": ["whatweb", "--version"],
        "apt_package": "whatweb",
        "category": "Reconnaissance & OSINT",
        "description": "Scanner d'empreintes technologiques Web",
    },
    "wafw00f": {
        "check_cmd": ["wafw00f", "--version"],
        "apt_package": "wafw00f",
        "category": "Reconnaissance & OSINT",
        "description": "Détection de pare-feux applicatifs web (WAF)",
    },
    "wpscan": {
        "check_cmd": ["wpscan", "--version"],
        "apt_package": "wpscan",
        "category": "Audit Web & CMS Spécifiques",
        "description": "Scanner spécialisé de vulnérabilités WordPress",
    },
    "gowitness": {
        "check_cmd": ["gowitness", "version"],
        "apt_package": "gowitness",
        "category": "Rendus Visuels, Rapports & Télémétrie",
        "description": "Outil de capture d'écran web en masse",
    },
    "netdiscover": {
        "check_cmd": ["netdiscover", "-h"],
        "apt_package": "netdiscover",
        "category": "Scans Réseau & Exploration LAN",
        "description": "Exploration ARP active/passive pour réseau local",
    },
    "nbtscan": {
        "check_cmd": ["nbtscan", "-h"],
        "apt_package": "nbtscan",
        "category": "Scans Réseau & Exploration LAN",
        "description": "Résolution de nom d'hôte et groupes de travail NetBIOS",
    },
    "enum4linux": {
        "check_cmd": ["enum4linux", "-h"],
        "apt_package": "enum4linux",
        "category": "Scans Réseau & Exploration LAN",
        "description": "Énumération complète des informations Windows et partages Samba/SMB",
    },
    "smbmap": {
        "check_cmd": ["smbmap", "-h"],
        "apt_package": "smbmap",
        "category": "Scans Réseau & Exploration LAN",
        "description": "Énumération des droits d'accès et partages réseau SMB",
    },
    "psutil": {
        "py_module": "psutil",
        "pip_package": "psutil",
        "category": "Rendus Visuels, Rapports & Télémétrie",
        "description": "Statistiques système et monitoring CPU/RAM",
    },
    "python-docx": {
        "py_module": "docx",
        "pip_package": "python-docx",
        "category": "Rendus Visuels, Rapports & Télémétrie",
        "description": "Génération et exportation de rapports Word (.docx)",
    },
    "whois": {
        "check_cmd": ["whois", "--version"],
        "apt_package": "whois",
        "category": "Reconnaissance & OSINT",
        "description": "Consultation des informations du registre de nom de domaine (WHOIS)",
    },
    "sublist3r": {
        "check_cmd": ["sublist3r", "-h"],
        "apt_package": "sublist3r",
        "category": "Reconnaissance & OSINT",
        "description": "Énumération OSINT passive de sous-domaines web",
    },
    "subfinder": {
        "check_cmd": ["subfinder", "-version"],
        "apt_package": "subfinder",
        "category": "Reconnaissance & OSINT",
        "description": "Découverte rapide multi-sources de sous-domaines web",
    },
    "theharvester": {
        "check_cmd": ["theHarvester", "-h"],
        "apt_package": "theharvester",
        "category": "Reconnaissance & OSINT",
        "description": "Collecte d'emails, sous-domaines et adresses IP (OSINT)",
    },
    "reportlab": {
        "py_module": "reportlab",
        "pip_package": "reportlab",
        "category": "Rendus Visuels, Rapports & Télémétrie",
        "description": "Génération et exportation de rapports PDF",
    },
    "joomscan": {
        "check_cmd": ["joomscan", "--help"],
        "apt_package": "joomscan",
        "category": "Audit Web & CMS Spécifiques",
        "description": "Scanner spécialisé de vulnérabilités pour le CMS Joomla",
    },
    "droopescan": {
        "py_module": "droopescan",
        "pip_package": "droopescan",
        "check_cmd": ["droopescan", "--help"],
        "apt_package": "droopescan",
        "category": "Audit Web & CMS Spécifiques",
        "description": "Scanner de vulnérabilités pour CMS Drupal et SilverStripe",
    },
    "moodlescan": {
        "check_cmd": ["moodlescan", "-h"],
        "apt_package": "moodlescan",
        "category": "Audit Web & CMS Spécifiques",
        "description": "Scanner spécialisé dans l'audit de sécurité des plateformes Moodle",
    },
    "paramspider": {
        "check_cmd": ["paramspider", "--help"],
        "apt_package": "paramspider",
        "category": "Reconnaissance & OSINT",
        "description": "Exploration des archives web pour extraire les URLs avec paramètres",
    },
    "sqlmap": {
        "check_cmd": ["sqlmap", "--version"],
        "apt_package": "sqlmap",
        "category": "Audit & Injections Bases de Données",
        "description": "Outil automatique de détection et d'exploitation d'injections SQL",
    },
    "nosqlmap": {
        "check_cmd": ["nosqlmap", "-h"],
        "apt_package": "nosqlmap",
        "category": "Audit & Injections Bases de Données",
        "description": "Outil d'audit d'injection et de sécurité pour bases de données NoSQL",
    },
    "gobuster": {
        "check_cmd": ["gobuster", "version"],
        "apt_package": "gobuster",
        "category": "Fuzzing & Recherche de Vulnérabilités",
        "description": "Fuzzer d'arborescence réseau, de répertoires web et de sous-domaines",
    },
}


# =====================================================================================
# SECTION 1 : VÉRIFICATION DES PRÉREQUIS
# =====================================================================================

def _tool_is_installed(cmd):
    """Vérifie qu'un binaire est présent et exécutable dans le PATH système ou le venv local."""
    binary = cmd[0]
    if shutil.which(binary) is not None:
        return True
    venv_bin = os.path.join(os.path.dirname(os.path.abspath(__file__)), "venv", "bin", binary)
    if os.path.isfile(venv_bin) and os.access(venv_bin, os.X_OK):
        return True
    return False

def _py_module_is_installed(mod_name):
    """Vérifie si un module Python est installé et importable."""
    try:
        return importlib.util.find_spec(mod_name) is not None
    except Exception:
        return False


@app.route("/api/check-deps", methods=["GET"])
def check_deps():
    """
    Renvoie l'état d'installation de chaque outil requis.
    Le frontend utilise cette route au chargement de la page pour afficher
    (ou non) le bandeau d'alerte "dépendances manquantes".
    """
    status = {}
    missing_packages = []
    for name, meta in REQUIRED_TOOLS.items():
        installed = False
        if "py_module" in meta:
            installed = _py_module_is_installed(meta["py_module"])
        if not installed and "check_cmd" in meta:
            installed = _tool_is_installed(meta["check_cmd"])

        pkg_name = meta.get("pip_package") or meta.get("apt_package", name)

        if not installed:
            missing_packages.append(pkg_name)

        status[name] = {
            "installed": installed,
            "description": meta["description"],
            "category": meta.get("category", "Divers / Système"),
            "is_python": "py_module" in meta,
        }
    return jsonify({
        "all_installed": len(missing_packages) == 0,
        "tools": status,
        "missing_apt_packages": missing_packages,
    })


LAST_NET_IO = {"bytes_recv": 0, "bytes_sent": 0, "time": 0}
PUBLIC_IP_CACHE = {"ip": "Détection...", "last_check": 0, "fetching": False}

def is_proxy_port_open(host, port):
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.settimeout(0.05)
        return s.connect_ex((host, port)) == 0


def _fetch_public_ip_worker():
    global PUBLIC_IP_CACHE
    services = [
        ("https://api.ipify.org?format=json", "json", "ip"),
        ("https://ifconfig.me/ip", "text", None),
        ("https://icanhazip.com", "text", None),
    ]
    detected_ip = None

    # 1. Tester les proxies SOCKS/HTTP OnionHop & Tor s'ils sont actifs
    proxy_candidates = [
        ("socks5h://127.0.0.1:9050", "127.0.0.1", 9050),
        ("socks5h://127.0.0.1:9150", "127.0.0.1", 9150),
        ("socks5h://127.0.0.1:1080", "127.0.0.1", 1080),
        ("socks5h://127.0.0.1:9052", "127.0.0.1", 9052),
    ]

    for env_var in ["ALL_PROXY", "HTTP_PROXY", "HTTPS_PROXY", "socks_proxy"]:
        val = os.environ.get(env_var)
        if val:
            try:
                parsed = urllib.parse.urlparse(val)
                if parsed.hostname and parsed.port:
                    proxy_candidates.insert(0, (val, parsed.hostname, parsed.port))
            except Exception:
                pass

    for proxy_str, host, port in proxy_candidates:
        if not is_proxy_port_open(host, port):
            continue
        for url, fmt, key in services:
            try:
                cmd = ["curl", "-s", "--max-time", "6", "-A", "Mozilla/5.0", "--proxy", proxy_str, url]
                res = subprocess.run(cmd, capture_output=True, text=True, timeout=7)
                if res.returncode == 0 and res.stdout:
                    raw = res.stdout.strip()
                    if fmt == "json":
                        data = json.loads(raw)
                        ip_candidate = data.get(key, "").strip()
                    else:
                        ip_candidate = raw.strip()
                    if ip_candidate:
                        ipaddress.ip_address(ip_candidate)
                        detected_ip = ip_candidate
                        break
            except Exception:
                continue
        if detected_ip:
            break

    # 2. Si aucun proxy n'est actif, passer en connexion directe
    if not detected_ip:
        for url, fmt, key in services:
            try:
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0 (X11; Linux x86_64)"})
                with urllib.request.urlopen(req, timeout=3) as resp:
                    if resp.status == 200:
                        raw = resp.read().decode("utf-8").strip()
                        if fmt == "json":
                            data = json.loads(raw)
                            ip_candidate = data.get(key, "").strip()
                        else:
                            ip_candidate = raw.strip()
                        if ip_candidate:
                            ipaddress.ip_address(ip_candidate)
                            detected_ip = ip_candidate
                            break
            except Exception:
                continue

    PUBLIC_IP_CACHE["ip"] = detected_ip if detected_ip else "Indisponible"
    PUBLIC_IP_CACHE["last_check"] = time.time()
    PUBLIC_IP_CACHE["fetching"] = False


def get_public_ip():
    global PUBLIC_IP_CACHE
    now = time.time()
    if (now - PUBLIC_IP_CACHE["last_check"] > 10 or PUBLIC_IP_CACHE["ip"] in ["Détection...", "Indisponible"]) and not PUBLIC_IP_CACHE["fetching"]:
        PUBLIC_IP_CACHE["fetching"] = True
        threading.Thread(target=_fetch_public_ip_worker, daemon=True).start()
    return PUBLIC_IP_CACHE["ip"]


@app.route("/api/system-stats")
def api_system_stats():
    global LAST_NET_IO
    now = time.time()
    cpu_percent = psutil.cpu_percent(interval=None)
    mem = psutil.virtual_memory()

    net_io = psutil.net_io_counters()
    rx_speed = 0.0
    tx_speed = 0.0

    if LAST_NET_IO["time"] > 0 and now > LAST_NET_IO["time"]:
        dt = now - LAST_NET_IO["time"]
        rx_speed = (net_io.bytes_recv - LAST_NET_IO["bytes_recv"]) / dt
        tx_speed = (net_io.bytes_sent - LAST_NET_IO["bytes_sent"]) / dt

    LAST_NET_IO["bytes_recv"] = net_io.bytes_recv
    LAST_NET_IO["bytes_sent"] = net_io.bytes_sent
    LAST_NET_IO["time"] = now

    public_ip = get_public_ip()

    return jsonify({
        "cpu_percent": round(cpu_percent, 1),
        "ram_percent": round(mem.percent, 1),
        "ram_used_gb": round(mem.used / (1024 ** 3), 2),
        "ram_total_gb": round(mem.total / (1024 ** 3), 2),
        "rx_bytes_sec": round(rx_speed, 1),
        "tx_bytes_sec": round(tx_speed, 1),
        "public_ip": public_ip,
    })


def _run_install_thread(packages):
    """
    Exécute `pkexec apt-get install -y <paquets>` dans un thread séparé.

    Pourquoi pkexec et pas sudo ?
    -----------------------------
    - `sudo` en ligne de commande exigerait de faire transiter le mot de passe
      root via le processus Python (mauvaise pratique, terminal bloquant).
    - `pkexec` délègue l'authentification à Polkit, qui affiche une fenêtre
      graphique native du système (GNOME/KDE) demandant le mot de passe.
      Le mot de passe ne transite JAMAIS par notre application : Polkit
      s'exécute en dehors du processus Flask.
    - Le thread Python ne fait qu'attendre la fin du processus pkexec, sans
      bloquer le reste du serveur Flask (les autres requêtes HTTP continuent
      d'être traitées).
    """
    INSTALL_STATUS["running"] = True
    INSTALL_STATUS["success"] = None
    INSTALL_STATUS["log"] = "Lancement de pkexec, veuillez valider la fenêtre système...\n"
    INSTALL_STATUS["finished_at"] = None

    cmd = ["pkexec", "apt-get", "install", "-y"] + packages
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=300,  # 5 minutes max
        )
        INSTALL_STATUS["log"] += result.stdout + "\n" + result.stderr

        # Codes de sortie usuels pour pkexec :
        #   0   -> succès
        #   126 -> authentification annulée par l'utilisateur (pop-up fermée)
        #   127 -> pkexec/commande introuvable
        if result.returncode == 0:
            INSTALL_STATUS["success"] = True
        elif result.returncode == 126:
            INSTALL_STATUS["success"] = False
            INSTALL_STATUS["log"] += "\n[Annulé par l'utilisateur : authentification refusée ou fermée]"
        else:
            INSTALL_STATUS["success"] = False
            INSTALL_STATUS["log"] += f"\n[Échec, code de retour {result.returncode}]"

    except FileNotFoundError:
        INSTALL_STATUS["success"] = False
        INSTALL_STATUS["log"] += (
            "\n[Erreur : la commande 'pkexec' est introuvable sur ce système. "
            "Installez le paquet 'policykit-1' ou lancez l'installation manuellement : "
            f"sudo apt-get install -y {' '.join(packages)}]"
        )
    except subprocess.TimeoutExpired:
        INSTALL_STATUS["success"] = False
        INSTALL_STATUS["log"] += "\n[Délai d'attente dépassé]"
    finally:
        INSTALL_STATUS["running"] = False
        INSTALL_STATUS["finished_at"] = datetime.now().isoformat()


@app.route("/api/install-deps", methods=["POST"])
def install_deps():
    """
    Déclenche l'installation des paquets manquants via pkexec.
    Non-bloquant : renvoie immédiatement, le frontend doit poller
    /api/install-status pour suivre la progression.
    """
    if INSTALL_STATUS["running"]:
        return jsonify({"error": "Une installation est déjà en cours."}), 409

    data = request.get_json(silent=True) or {}
    packages = data.get("packages")

    # Si aucun paquet n'est explicitement fourni, on recalcule la liste
    # des paquets manquants côté serveur (source de vérité)
    if not packages:
        packages = []
        for name, meta in REQUIRED_TOOLS.items():
            if "apt_package" in meta and "check_cmd" in meta:
                if not _tool_is_installed(meta["check_cmd"]):
                    packages.append(meta["apt_package"])

    # Liste blanche stricte des paquets apt
    allowed_packages = {meta["apt_package"] for meta in REQUIRED_TOOLS.values() if "apt_package" in meta}
    packages = [p for p in packages if p in allowed_packages]

    if not packages:
        return jsonify({"error": "Aucun paquet valide à installer."}), 400

    thread = threading.Thread(target=_run_install_thread, args=(packages,), daemon=True)
    thread.start()

    return jsonify({"started": True, "packages": packages})


@app.route("/api/install-status", methods=["GET"])
def install_status():
    """Route de polling utilisée par le frontend pendant l'installation."""
    return jsonify(INSTALL_STATUS)


# =====================================================================================
# SECTION 2 : VALIDATION DES ENTRÉES UTILISATEUR (sécurité)
# =====================================================================================

HOSTNAME_RE = re.compile(
    r"^(?=.{1,253}$)(?!-)[A-Za-z0-9-]{1,63}(?<!-)"
    r"(\.(?!-)[A-Za-z0-9-]{1,63}(?<!-))*$"
)


def validate_target(target):
    """
    Valide et parse la cible (IP, CIDR, hostname, ou plages d'IP multiples).
    Supporte les virgules et les plages du type 192.168.1.11-192.168.1.21.
    """
    target = target.strip()
    if not target:
        return False, "La cible ne peut pas être vide."

    # Séparation par virgules et espaces pour gérer les cibles multiples
    raw_targets = [t.strip() for t in target.replace(',', ' ').split() if t.strip()]
    if not raw_targets:
        return False, "Cible invalide."

    valid_targets = []
    
    for t in raw_targets:
        # Si la cible est une URL (http/https), extraire l'hôte
        if t.startswith("http://") or t.startswith("https://"):
            parsed = urlparse(t)
            t = parsed.netloc.split(":")[0] if parsed.netloc else parsed.path.split("/")[0].split(":")[0]

        # Vérification des plages d'IP (ex: 192.168.1.11-192.168.1.21 ou 192.168.1.11-21)
        if '-' in t:
            parts = t.split('-')
            if len(parts) == 2:
                start_ip, end_ip = parts[0].strip(), parts[1].strip()
                # Conversion du format IP-IP complet vers IP-fin attendu par Nmap
                if '.' in end_ip:
                    start_octets = start_ip.split('.')
                    end_octets = end_ip.split('.')
                    if len(start_octets) == 4 and len(end_octets) == 4:
                        if start_octets[:3] == end_octets[:3]:
                            t = f"{start_ip}-{end_octets[3]}"
                        else:
                            return False, f"Plage invalide (les IP doivent être dans le même /24) : {t}"
                
                # Vérification de sécurité de la chaîne finale pour Nmap
                if re.match(r'^[\w\.\-]+$', t):
                    valid_targets.append(t)
                    continue

        # IP simple ou réseau CIDR
        try:
            ipaddress.ip_network(t, strict=False)
            valid_targets.append(t)
            continue
        except ValueError:
            pass

        # Nom d'hôte (ex: localhost, machine.local)
        if HOSTNAME_RE.match(t):
            valid_targets.append(t)
            continue

        return False, f"Cible ou format invalide : {t}"

    return True, valid_targets


ALLOWED_TIMING = {"0", "1", "2", "3", "4", "5"}
ALLOWED_SCAN_TYPES = {"sS", "sT", "sU", "sA", "sn"}
ALLOWED_PORT_MODES = {"fast", "top1000", "top_1000", "all", "custom"}
CUSTOM_PORTS_RE = re.compile(r"^[0-9,\-]{1,200}$")


# =====================================================================================
# SECTION 3 : CONSTRUCTION DE LA COMMANDE NMAP À PARTIR DES OPTIONS GRAPHIQUES
# =====================================================================================

def build_nmap_command(options):
    """
    Traduit les options choisies graphiquement par l'élève en flags nmap.
    Chaque flag est documenté ici et dans l'interface (voir index.html)
    pour que l'élève comprenne ce qu'il déclenche réellement.
    """
    cmd = ["nmap"]
    explanation = []  # trace pédagogique des flags utilisés, renvoyée au frontend

    scan_type = options.get("scan_type", "sS")
    if scan_type not in ALLOWED_SCAN_TYPES:
        scan_type = "sS"
    cmd.append(f"-{scan_type}")
    explanation.append({
        "flag": f"-{scan_type}",
        "meaning": {
            "sS": "Scan SYN furtif (half-open) : rapide, ne complète pas la connexion TCP.",
            "sT": "Scan Connect() complet : plus lent mais ne nécessite pas les privilèges root.",
            "sU": "Scan UDP : plus lent, détecte les services UDP (DNS, SNMP...).",
            "sA": "Scan ACK : sert à cartographier les règles de pare-feu (stateful vs stateless).",
            "sn": "Ping scan uniquement : détecte les hôtes en vie sans scanner les ports.",
        }[scan_type],
    })

    # Détection de version des services (-sV)
    if options.get("service_version"):
        cmd.append("-sV")
        explanation.append({
            "flag": "-sV",
            "meaning": "Sonde chaque port ouvert pour identifier le service et sa version exacte (ex: Apache 2.4.29).",
        })

    # Détection d'OS (-O) - nécessite les privilèges root, on prévient l'élève
    if options.get("os_detection"):
        cmd.append("-O")
        explanation.append({
            "flag": "-O",
            "meaning": "Tente de deviner le système d'exploitation de la cible via l'empreinte de sa pile TCP/IP.",
        })

    # Scripts NSE par défaut (--script=default) : detection de vulnérabilités basiques
    if options.get("default_scripts"):
        cmd.append("--script=default")
        explanation.append({
            "flag": "--script=default",
            "meaning": "Exécute les scripts NSE 'sûrs' par défaut (bannières, infos basiques, quelques vulnérabilités connues).",
        })

    if options.get("vuln_scripts"):
        cmd.append("--script=vuln")
        explanation.append({
            "flag": "--script=vuln",
            "meaning": "Exécute les scripts NSE dédiés à la détection de vulnérabilités connues (plus long, plus intrusif).",
        })

    if options.get("vulners_script"):
        cmd.append("--script=vulners")
        explanation.append({
            "flag": "--script=vulners",
            "meaning": "Interroge l'API Vulners avec les versions détectées pour lister les CVE et scores CVSS associés.",
        })

    if options.get("aggressive_scan"):
        cmd.append("-A")
        explanation.append({
            "flag": "-A",
            "meaning": "Scan agressif : active la détection d'OS, la détection de version, le scan de scripts (-sC) et traceroute. Nécessite les privilèges root.",
        })

    # Ports à scanner
    port_mode = options.get("port_mode", "top1000")
    if port_mode not in ALLOWED_PORT_MODES:
        port_mode = "top1000"

    if port_mode == "fast":
        cmd.append("-F")
        explanation.append({"flag": "-F", "meaning": "Scan rapide : seulement les 100 ports les plus courants."})
    elif port_mode in ["top1000", "top_1000"]:
        # comportement par défaut de nmap (top 1000 ports), pas de flag nécessaire
        explanation.append({"flag": "(défaut)", "meaning": "Scan des 1000 ports les plus courants (comportement par défaut de nmap)."})
    elif port_mode == "all":
        cmd.append("-p-")
        explanation.append({"flag": "-p-", "meaning": "Scan de l'intégralité des 65535 ports TCP (le plus long, le plus exhaustif)."})
    elif port_mode == "custom":
        custom_ports = str(options.get("custom_ports", ""))
        if CUSTOM_PORTS_RE.match(custom_ports):
            cmd.extend(["-p", custom_ports])
            explanation.append({"flag": f"-p {custom_ports}", "meaning": "Scan restreint à la liste/plage de ports spécifiée."})

    # Timing template (intensité / vitesse du scan, slider -T0 à -T5)
    timing = str(options.get("timing", "3"))
    if timing not in ALLOWED_TIMING:
        timing = "3"
    cmd.append(f"-T{timing}")
    timing_meanings = {
        "0": "Paranoïaque : extrêmement lent, pour passer sous le radar des IDS.",
        "1": "Furtif : très lent, discret.",
        "2": "Poli : réduit la charge réseau, plus lent que la normale.",
        "3": "Normal : vitesse par défaut de nmap.",
        "4": "Agressif : rapide, suppose un réseau fiable et rapide.",
        "5": "Insane : le plus rapide possible, peut manquer des résultats sur réseau lent.",
    }
    explanation.append({"flag": f"-T{timing}", "meaning": timing_meanings[timing]})

    if options.get("skip_host_discovery") or options.get("scan_type") == "sS" or options.get("is_local"):
        cmd.append("-Pn")
        explanation.append({
            "flag": "-Pn",
            "meaning": "Désactive le ping préalable : Nmap scanne directement les ports sans vérifier d'abord si l’hôte répond aux ICMP/Echo.",
        })

    # Optimisations de vitesse & performance (évite les blocages sur cibles filtrées)
    max_retries = str(options.get("max_retries", "2"))
    cmd.extend(["--max-retries", max_retries])
    explanation.append({
        "flag": f"--max-retries {max_retries}",
        "meaning": f"Limite le nombre de tentatives de retransmission à {max_retries} par port (accélère le traitement des ports filtrés).",
    })

    host_timeout = str(options.get("host_timeout", "5m"))
    cmd.extend(["--host-timeout", host_timeout])
    explanation.append({
        "flag": f"--host-timeout {host_timeout}",
        "meaning": f"Délai d'attente maximum global accordé à l'hôte ({host_timeout}).",
    })

    if options.get("min_rate"):
        min_rate = str(options.get("min_rate"))
        cmd.extend(["--min-rate", min_rate])
        explanation.append({
            "flag": f"--min-rate {min_rate}",
            "meaning": f"Impose un débit d'émission minimal de {min_rate} paquets/seconde.",
        })

    # Mode verbeux + rapport de statistiques périodique.
    cmd.append("-v")
    cmd.extend(["--stats-every", "2s"])
    explanation.append({
        "flag": "-v --stats-every 2s",
        "meaning": "Active l'affichage de la progression du scan toutes les 2 secondes (visible dans le terminal de suivi).",
    })

    # Sortie XML sur stdout pour parsing structuré
    cmd.extend(["-oX", "-"])

    return cmd, explanation


def resolve_host_identity(ip, mac=None, vendor=None):
    """
    Enrichit les informations d'un équipement local en interrogeant NetBIOS (nbtscan),
    le DNS inversé et en analysant le préfixe MAC / Constructeur.
    """
    identity = {
        "hostname": "",
        "netbios_name": "",
        "workgroup": "",
        "vendor": vendor or "Inconnu",
        "category": "Équipement Réseau",
        "icon": "fa-server"
    }

    if not ip or ip == "inconnu":
        return identity

    # 1. Résolution DNS inversée silencieuse avec timeout très court (0.5s)
    try:
        orig_timeout = socket.getdefaulttimeout()
        socket.setdefaulttimeout(0.5)
        host_tuple = socket.gethostbyaddr(ip)
        socket.setdefaulttimeout(orig_timeout)
        if host_tuple and host_tuple[0]:
            identity["hostname"] = host_tuple[0]
    except Exception:
        pass

    # 2. Résolution NetBIOS via nbtscan (timeout 0.8s)
    if shutil.which("nbtscan"):
        try:
            res = subprocess.run(
                ["nbtscan", "-m", ip],
                capture_output=True,
                text=True,
                timeout=0.8
            )
            if res.stdout:
                lines = res.stdout.strip().split("\n")
                for line in lines:
                    if ip in line:
                        parts = line.split()
                        if len(parts) >= 2:
                            identity["netbios_name"] = parts[1]
                        if len(parts) >= 4:
                            identity["workgroup"] = parts[3]
                        break
        except Exception:
            pass

    # 3. Dictionnaires de marques / catégories basés sur le nom, le vendeur MAC ou le NetBIOS
    search_str = f"{identity['vendor']} {identity['hostname']} {identity['netbios_name']}".lower()

    if any(k in search_str for k in ["freebox", "free sas", "freebox2", "livebox", "sagemcom", "technicolor", "netgear", "tp-link", "cisco", "asus", "huawei", "zte"]):
        identity["category"] = "Routeur / Répéteur / Box"
        identity["icon"] = "fa-wifi"
    elif any(k in search_str for k in ["apple", "iphone", "ipad", "macbook", "imac", "apple-tv"]):
        identity["category"] = "Équipement Apple"
        identity["icon"] = "fa-apple"
    elif any(k in search_str for k in ["samsung", "lg", "philips", "sony", "tcl", "chromecast", "firetv", "tv"]):
        identity["category"] = "Smart TV / Médias"
        identity["icon"] = "fa-tv"
    elif any(k in search_str for k in ["hp", "hewlett", "epson", "canon", "brother", "xerox", "printer"]):
        identity["category"] = "Imprimante Réseau"
        identity["icon"] = "fa-print"
    elif any(k in search_str for k in ["synology", "qnap", "nas", "truenas"]):
        identity["category"] = "Serveur NAS / Stockage"
        identity["icon"] = "fa-database"
    elif any(k in search_str for k in ["espressif", "tuya", "sonoff", "shelly", "raspberry", "arduino", "esp8266", "esp32"]):
        identity["category"] = "Objet Connecté / IoT"
        identity["icon"] = "fa-microchip"
    elif any(k in search_str for k in ["windows", "desktop", "workstation", "pc", "dell", "lenovo", "acer", "msi"]):
        identity["category"] = "Station PC / Windows"
        identity["icon"] = "fa-desktop"

    return identity


# =====================================================================================
# SECTION 4 : EXÉCUTION DU SCAN ET PARSING XML
# =====================================================================================

def parse_nmap_xml(xml_output):
    """Parse la sortie XML native de nmap en structure JSON exploitable par le frontend."""
    root = ET.fromstring(xml_output)
    hosts_result = []

    for host in root.findall("host"):
        status_el = host.find("status")
        host_state = status_el.get("state") if status_el is not None else "up"

        ip_addr = "inconnu"
        mac_addr = None
        vendor = None
        for addr_el in host.findall("address"):
            addr_type = addr_el.get("addrtype")
            if addr_type in ["ipv4", "ipv6"]:
                ip_addr = addr_el.get("addr")
            elif addr_type == "mac":
                mac_addr = addr_el.get("addr")
                vendor = addr_el.get("vendor")

        hostname_el = host.find("hostnames/hostname")
        hostname = hostname_el.get("name") if hostname_el is not None else None

        # OS detection
        os_matches = []
        os_el = host.find("os")
        if os_el is not None:
            for match in os_el.findall("osmatch"):
                os_matches.append({
                    "name": match.get("name"),
                    "accuracy": match.get("accuracy"),
                })

        ports_result = []
        ports_el = host.find("ports")
        if ports_el is not None:
            for port in ports_el.findall("port"):
                state_el = port.find("state")
                service_el = port.find("service")

                scripts = []
                for script in port.findall("script"):
                    scripts.append({
                        "id": script.get("id"),
                        "output": script.get("output", "")[:2000],  # borne la taille
                    })

                ports_result.append({
                    "port": port.get("portid"),
                    "protocol": port.get("protocol"),
                    "state": state_el.get("state") if state_el is not None else "unknown",
                    "service": service_el.get("name") if service_el is not None else "",
                    "product": service_el.get("product") if service_el is not None else "",
                    "version": service_el.get("version") if service_el is not None else "",
                    "scripts": scripts,
                })

        identity = resolve_host_identity(ip_addr, mac_addr, vendor)

        hosts_result.append({
            "ip": ip_addr,
            "mac": mac_addr,
            "vendor": identity["vendor"],
            "hostname": hostname or identity["hostname"] or identity["netbios_name"],
            "netbios_name": identity["netbios_name"],
            "workgroup": identity["workgroup"],
            "category": identity["category"],
            "icon": identity["icon"],
            "status": host_state,
            "os_matches": os_matches,
            "ports": ports_result,
        })

    return hosts_result


MAX_LOG_LINES = 5000  # borne de sécurité pour éviter une consommation mémoire illimitée


def _run_scan_thread(job_id, cmd, target, mode="distant"):
    """
    Lance nmap avec subprocess.Popen (au lieu de subprocess.run) afin de pouvoir
    lire sa sortie AU FUR ET À MESURE, plutôt que d'attendre la fin complète du
    processus. Deux threads lecteurs dédiés évitent les blocages classiques liés
    aux tubes (pipes) : un pipe non lu qui se remplit peut bloquer le processus
    enfant indéfiniment (deadlock) si stdout et stderr ne sont pas consommés en
    parallèle.

        - stdout -> accumulé silencieusement, c'est le flux XML final (-oX -),
                    parsé uniquement une fois le scan terminé.
        - stderr -> chaque ligne (messages -v / --stats-every) est ajoutée en
                    direct à job["log"], relu par le frontend via polling sur
                    /api/scan-status/<job_id> pour afficher le "terminal live".
    """
    SCAN_JOBS[job_id]["status"] = "running"
    SCAN_JOBS[job_id]["log"] = []

    # stdbuf force un buffering ligne par ligne côté nmap (au lieu du buffering
    # par bloc utilisé automatiquement quand la sortie n'est pas un vrai
    # terminal), pour que les lignes de progression arrivent sans latence.
    run_cmd = cmd
    input_data = None
    if CONFIGURED_SUDO_PASSWORD:
        run_cmd = ["sudo", "-S"] + cmd
        input_data = f"{CONFIGURED_SUDO_PASSWORD}\n"
    elif shutil.which("stdbuf"):
        run_cmd = ["stdbuf", "-oL", "-eL"] + cmd

    run_cmd = _wrap_with_proxy(run_cmd, use_sudo=bool(CONFIGURED_SUDO_PASSWORD))

    def append_log(line):
        log = SCAN_JOBS[job_id]["log"]
        log.append(line)
        if len(log) > MAX_LOG_LINES:
            del log[0:len(log) - MAX_LOG_LINES]

    try:
        process = subprocess.Popen(
            run_cmd,
            stdin=subprocess.PIPE if input_data else None,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
        )
        if input_data:
            try:
                process.stdin.write(input_data)
                process.stdin.flush()
            except Exception:
                pass
        ACTIVE_PROCESSES[job_id] = process
    except FileNotFoundError:
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = "La commande nmap est introuvable."
        return
    except PermissionError:
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = (
            "Permission refusée : certains types de scans (-sS, -O) nécessitent "
            "les privilèges root. Lancez le serveur Flask avec les capacités "
            "réseau adéquates (voir README, section setcap)."
        )
        return

    stdout_chunks = []

    def read_stdout():
        for chunk in process.stdout:
            stdout_chunks.append(chunk)

    def read_stderr():
        for line in process.stderr:
            clean_line = line.rstrip("\n")
            if clean_line:
                append_log(clean_line)

    t_out = threading.Thread(target=read_stdout, daemon=True)
    t_err = threading.Thread(target=read_stderr, daemon=True)
    t_out.start()
    t_err.start()

    try:
        process.wait(timeout=900)  # 15 minutes max, un scan -p- peut être long
    except subprocess.TimeoutExpired:
        process.kill()
        process.wait()
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = "Délai d'attente dépassé (15 minutes). Réduisez la portée du scan."
        append_log("[✘ Scan interrompu : délai dépassé]")
        return

    t_out.join(timeout=5)
    t_err.join(timeout=5)

    full_stdout = "".join(stdout_chunks)

    if process.returncode != 0 and not full_stdout.strip():
        err_log = "\n".join(SCAN_JOBS[job_id]["log"])
        if "-sS" in cmd and ("root" in err_log.lower() or "privilèg" in err_log.lower()):
            append_log("[!] Note : Le scan SYN (-sS) nécessite les privilèges root. Basculement automatique en mode TCP Connect (-sT)...")
            fallback_cmd = [c if c != "-sS" else "-sT" for c in cmd]
            try:
                fb_proc = subprocess.run(fallback_cmd, capture_output=True, text=True, timeout=600)
                if fb_proc.returncode == 0 and fb_proc.stdout.strip():
                    full_stdout = fb_proc.stdout
                else:
                    SCAN_JOBS[job_id]["status"] = "error"
                    SCAN_JOBS[job_id]["error"] = fb_proc.stderr or "Échec du scan nmap fallback."
                    return
            except Exception as e:
                SCAN_JOBS[job_id]["status"] = "error"
                SCAN_JOBS[job_id]["error"] = f"Erreur lors du fallback nmap : {e}"
                return
        else:
            SCAN_JOBS[job_id]["status"] = "error"
            SCAN_JOBS[job_id]["error"] = (
                "\n".join(SCAN_JOBS[job_id]["log"][-15:]) or "nmap a échoué sans sortie."
            )
            return

    try:
        parsed = parse_nmap_xml(full_stdout)
    except ET.ParseError:
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = "Impossible d'analyser la sortie XML de nmap."
        return
    except Exception as exc:  # sécurité : ne jamais laisser planter le thread silencieusement
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = f"Erreur inattendue lors du parsing : {exc}"
        return

    SCAN_JOBS[job_id]["status"] = "done"
    save_history("nmap", target, parsed, mode=mode)
    SCAN_JOBS[job_id]["result"] = parsed
    append_log("[✔ Scan terminé]")


@app.route("/api/scan", methods=["POST"])
def start_scan():
    """
    Lance un scan nmap asynchrone à partir des options graphiques envoyées
    par le frontend. Renvoie un job_id à poller via /api/scan-status/<id>.
    """
    data = request.get_json(silent=True) or {}
    target = data.get("target", "")
    mode = data.get("mode", "distant")

    is_valid, msg = validate_target(target)
    if not is_valid:
        return jsonify({"error": msg}), 400

    # Vérification que nmap est bien installé avant de lancer quoi que ce soit
    if not _tool_is_installed(REQUIRED_TOOLS["nmap"]["check_cmd"]):
        return jsonify({"error": "nmap n'est pas installé. Utilisez le bouton d'installation des dépendances."}), 412

    cmd, explanation = build_nmap_command(data.get("options", {}))
    cmd.extend(msg)  # liste de cibles validées ajoutées à la fin de la commande

    job_id = str(uuid.uuid4())
    SCAN_JOBS[job_id] = {
        "status": "queued",
        "result": None,
        "error": None,
        "log": [],
        "command": " ".join(cmd),
        "explanation": explanation,
        "started_at": datetime.now().isoformat(),
    }

    target_str = ", ".join(msg)
    thread = threading.Thread(target=_run_scan_thread, args=(job_id, cmd, target_str, mode), daemon=True)
    thread.start()

    return jsonify({
        "job_id": job_id,
        "command_preview": " ".join(cmd),
        "explanation": explanation,
    })


@app.route("/api/scan-status/<job_id>", methods=["GET"])
def scan_status(job_id):
    job = SCAN_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)


@app.route("/api/scan/stop", methods=["POST"])
def stop_all_scans():
    """
    Arrête immédiatement tous les scans Nmap, Netdiscover, Nuclei, WhatWeb en cours.
    """
    stopped_count = 0
    for job_id, proc in list(ACTIVE_PROCESSES.items()):
        if proc:
            try:
                proc.terminate()
                stopped_count += 1
            except Exception:
                pass
    ACTIVE_PROCESSES.clear()

    all_job_dicts = [SCAN_JOBS, NETDISCOVER_JOBS, NUCLEI_JOBS, WHATWEB_JOBS, WPSCAN_JOBS]
    for jobs in all_job_dicts:
        for job_id, job in list(jobs.items()):
            if job.get("status") in ["running", "queued"]:
                job["status"] = "error"
                job["error"] = "Scan interrompu par l'utilisateur."

    return jsonify({"success": True, "stopped_count": stopped_count})


# =====================================================================================
# SECTION 5 : RECHERCHE DE CVE / EXPLOITS VIA SEARCHSPLOIT
# =====================================================================================

SEARCHSPLOIT_CACHE = {}

@app.route("/api/cve-search", methods=["POST"])
def cve_search():
    """
    Recherche des exploits/CVE connus pour un couple (produit, version)
    détecté par nmap -sV, via la base Exploit-DB locale (searchsploit).
    Les résultats sont mis en cache mémoire pour accélérer les requêtes réseau.
    """
    data = request.get_json(silent=True) or {}
    product = str(data.get("product", "")).strip()
    version = str(data.get("version", "")).strip()

    if not product:
        return jsonify({"error": "Produit manquant."}), 400

    query = f"{product} {version}".strip()

    if query in SEARCHSPLOIT_CACHE:
        return jsonify({"query": query, "results": SEARCHSPLOIT_CACHE[query]})

    if not _tool_is_installed(REQUIRED_TOOLS["searchsploit"]["check_cmd"]):
        return jsonify({"error": "searchsploit n'est pas installé.", "results": []}), 412

    try:
        result = subprocess.run(
            ["searchsploit", "--json", query],
            capture_output=True,
            text=True,
            timeout=30,
        )
        results = []
        if result.stdout:
            try:
                payload = json.loads(result.stdout)
                for entry in payload.get("RESULTS_EXPLOIT", []):
                    results.append({
                        "title": entry.get("Title"),
                        "edb_id": entry.get("EDB-ID"),
                        "date": entry.get("Date"),
                        "type": entry.get("Type"),
                        "platform": entry.get("Platform"),
                    })
            except json.JSONDecodeError:
                pass

        # Si aucun résultat avec produit + version, réessayer avec produit uniquement
        if not results and version:
            res_prod = subprocess.run(
                ["searchsploit", "--json", product],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if res_prod.stdout:
                try:
                    payload = json.loads(res_prod.stdout)
                    for entry in payload.get("RESULTS_EXPLOIT", []):
                        results.append({
                            "title": entry.get("Title"),
                            "edb_id": entry.get("EDB-ID"),
                            "date": entry.get("Date"),
                            "type": entry.get("Type"),
                            "platform": entry.get("Platform"),
                        })
                except json.JSONDecodeError:
                    pass

        SEARCHSPLOIT_CACHE[query] = results
        return jsonify({"query": query, "results": results})

    except subprocess.TimeoutExpired:
        return jsonify({"error": "searchsploit a mis trop de temps à répondre."}), 504
    except FileNotFoundError:
        return jsonify({"error": "searchsploit introuvable."}), 412


# =====================================================================================
# SECTION 6 : NUCLEI SCANNER
# =====================================================================================

def _run_nuclei_thread(job_id, target_url, tags="", options=None):
    if options is None:
        options = {}
        
    NUCLEI_JOBS[job_id]["status"] = "running"
    NUCLEI_JOBS[job_id]["log"] = []
    
    cmd = ["nuclei", "-u", target_url, "-jsonl", "-nc"]
    if tags:
        cmd.extend(["-tags", tags])
        
    rate_limit = options.get("rate_limit")
    if rate_limit:
        cmd.extend(["-rl", str(rate_limit)])
    
    cmd = _wrap_with_proxy(cmd)
    
    def append_log(line):
        log = NUCLEI_JOBS[job_id]["log"]
        log.append(line)
        if len(log) > MAX_LOG_LINES:
            del log[0:len(log) - MAX_LOG_LINES]

    try:
        process = subprocess.Popen(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
        )
    except FileNotFoundError:
        NUCLEI_JOBS[job_id]["status"] = "error"
        NUCLEI_JOBS[job_id]["error"] = "La commande nuclei est introuvable."
        return

    stdout_chunks = []

    def read_stdout():
        for chunk in process.stdout:
            stdout_chunks.append(chunk)

    def read_stderr():
        for line in process.stderr:
            clean_line = line.rstrip("\n")
            if clean_line:
                append_log(clean_line)

    t_out = threading.Thread(target=read_stdout, daemon=True)
    t_err = threading.Thread(target=read_stderr, daemon=True)
    t_out.start()
    t_err.start()

    try:
        process.wait(timeout=1800)  # 30 minutes max
    except subprocess.TimeoutExpired:
        process.kill()
        process.wait()
        NUCLEI_JOBS[job_id]["status"] = "error"
        NUCLEI_JOBS[job_id]["error"] = "Délai d'attente dépassé (30 minutes)."
        append_log("[✘ Scan interrompu : délai dépassé]")
        return

    t_out.join(timeout=5)
    t_err.join(timeout=5)

    full_stdout = "".join(stdout_chunks)
    
    results = []
    for line in full_stdout.strip().split("\n"):
        if not line.strip():
            continue
        try:
            results.append(json.loads(line))
        except json.JSONDecodeError:
            pass

    NUCLEI_JOBS[job_id]["status"] = "done"
    NUCLEI_JOBS[job_id]["result"] = results
    save_history("nuclei", target_url, results)
    append_log("[✔ Scan terminé]")


@app.route("/api/nuclei-scan", methods=["POST"])
def start_nuclei_scan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()
    tags = data.get("tags", "").strip()

    options = data.get("options", {})

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["nuclei"]["check_cmd"]):
        return jsonify({"error": "nuclei n'est pas installé."}), 412

    job_id = str(uuid.uuid4())
    NUCLEI_JOBS[job_id] = {
        "status": "queued",
        "result": None,
        "error": None,
        "log": [],
        "target_url": target_url,
        "tags": tags,
        "started_at": datetime.now().isoformat(),
    }

    thread = threading.Thread(target=_run_nuclei_thread, args=(job_id, target_url, tags, options), daemon=True)
    thread.start()

    return jsonify({
        "job_id": job_id
    })

@app.route("/api/nuclei-status/<job_id>", methods=["GET"])
def nuclei_status(job_id):
    job = NUCLEI_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)


# =====================================================================================
# SECTION 7 : WHATWEB SCANNER
# =====================================================================================

def enrich_whatweb_with_cves(results):
    if not results or not isinstance(results, list):
        return results

    has_searchsploit = _tool_is_installed(REQUIRED_TOOLS["searchsploit"]["check_cmd"])

    for target_item in results:
        plugins = target_item.get("plugins", {})
        cve_list = []

        for plugin_name, plugin_data in plugins.items():
            name_lower = plugin_name.lower()
            if name_lower in ["country", "ip", "title", "script", "frame", "html5", "open-graph-protocol", "uncommonheaders", "strict-transport-security", "x-frame-options"]:
                continue

            versions = plugin_data.get("version", []) if isinstance(plugin_data, dict) else []
            version_str = versions[0] if versions else ""
            query = f"{plugin_name} {version_str}".strip()

            if has_searchsploit:
                try:
                    res_sp = subprocess.run(
                        ["searchsploit", "--json", query],
                        capture_output=True,
                        text=True,
                        timeout=8
                    )
                    if res_sp.stdout:
                        try:
                            payload = json.loads(res_sp.stdout)
                            exploits = payload.get("RESULTS_EXPLOIT", [])
                            for exp in exploits[:3]:
                                exp_type = str(exp.get("Type", ""))
                                is_remote = "remote" in exp_type.lower() or "webapps" in exp_type.lower()
                                cve_list.append({
                                    "plugin": plugin_name,
                                    "version": version_str,
                                    "title": exp.get("Title"),
                                    "edb_id": exp.get("EDB-ID"),
                                    "date": exp.get("Date"),
                                    "type": exp.get("Type"),
                                    "platform": exp.get("Platform"),
                                    "cvss_score": "7.5" if is_remote else "5.3",
                                    "severity": "Élevé" if is_remote else "Moyen"
                                })
                        except Exception:
                            pass
                except Exception:
                    pass

        target_item["cves_found"] = cve_list

    return results

def _run_whatweb_thread(job_id, target_url, options):
    WHATWEB_JOBS[job_id]["status"] = "running"
    
    with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as tmp_file:
        tmp_path = tmp_file.name

    cmd = ["whatweb"]
    cmd += ["-a", str(options.get("aggression", 1))]
    cmd += ["--follow-redirect", options.get("follow_redirect", "always")]
    cmd += ["--open-timeout", str(options.get("open_timeout", 15))]
    cmd += ["--read-timeout", str(options.get("read_timeout", 30))]

    user_agent = options.get("user_agent", "")
    if user_agent:
        cmd += ["-U", user_agent]

    if options.get("verbose"):
        cmd += ["-v"]

    if options.get("no_errors"):
        cmd += ["--no-errors"]

    if options.get("no_cookies"):
        cmd += ["--no-cookies"]

    wait_time = options.get("wait")
    if wait_time and int(wait_time) > 0:
        cmd += ["--wait", str(wait_time)]

    cmd += ["--log-json", tmp_path, "--colour=never", target_url]
    WHATWEB_JOBS[job_id]["command"] = " ".join(cmd)
    
    cmd = _wrap_with_proxy(cmd)

    try:
        process = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=max(180, options.get("open_timeout", 15) + options.get("read_timeout", 30) + 30),
        )
        
        if process.returncode != 0 and process.returncode != 1:
            WHATWEB_JOBS[job_id]["status"] = "error"
            WHATWEB_JOBS[job_id]["error"] = f"WhatWeb a échoué (code {process.returncode}): {process.stderr.strip() or process.stdout.strip()}"
            return

        results = []
        try:
            with open(tmp_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    results = json.loads(content)
        except json.JSONDecodeError as e:
            WHATWEB_JOBS[job_id]["status"] = "error"
            WHATWEB_JOBS[job_id]["error"] = f"Impossible de lire la sortie JSON de WhatWeb : {e}"
            return
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        if not results:
            full_output = process.stdout + "\n" + process.stderr
            if "ERROR Opening" in full_output or "ERROR:" in full_output:
                error_lines = [line.strip() for line in full_output.splitlines() if "ERROR" in line]
                err_msg = error_lines[0] if error_lines else "Erreur de connexion (cible injoignable ou bloquée)."
                WHATWEB_JOBS[job_id]["status"] = "error"
                WHATWEB_JOBS[job_id]["error"] = err_msg
                return

        # Enrichissement automatique avec les CVE & CVSS Searchsploit
        results = enrich_whatweb_with_cves(results)

        WHATWEB_JOBS[job_id]["status"] = "done"
        WHATWEB_JOBS[job_id]["result"] = results
        save_history("whatweb", target_url, results)

    except subprocess.TimeoutExpired:
        WHATWEB_JOBS[job_id]["status"] = "error"
        WHATWEB_JOBS[job_id]["error"] = "Délai d'attente dépassé pour WhatWeb."
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    except FileNotFoundError:
        WHATWEB_JOBS[job_id]["status"] = "error"
        WHATWEB_JOBS[job_id]["error"] = "La commande whatweb est introuvable."
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    except Exception as exc:
        WHATWEB_JOBS[job_id]["status"] = "error"
        WHATWEB_JOBS[job_id]["error"] = f"Erreur inattendue : {exc}"
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

@app.route("/api/whatweb-scan", methods=["POST"])
def start_whatweb_scan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()
    options = data.get("options", {})

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    # Validation de l'agressivité
    aggression = int(options.get("aggression", 1))
    if aggression not in [1, 3, 4]:
        return jsonify({"error": "Niveau d'agressivité invalide (valeurs acceptées : 1, 3, 4)."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["whatweb"]["check_cmd"]):
        return jsonify({"error": "whatweb n'est pas installé."}), 412

    job_id = str(uuid.uuid4())
    WHATWEB_JOBS[job_id] = {
        "status": "queued",
        "result": None,
        "error": None,
        "command": None,
        "target_url": target_url,
        "started_at": datetime.now().isoformat(),
    }

    thread = threading.Thread(target=_run_whatweb_thread, args=(job_id, target_url, options), daemon=True)
    thread.start()

    return jsonify({"job_id": job_id})

@app.route("/api/whatweb-status/<job_id>", methods=["GET"])
def whatweb_status(job_id):
    job = WHATWEB_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)

# =====================================================================================
# SECTION 8 : SCREENSHOT (GOWITNESS)
# =====================================================================================

SCREENSHOT_DIR = os.path.join(app.static_folder, 'screenshots')
os.makedirs(SCREENSHOT_DIR, exist_ok=True)

@app.route("/api/screenshot", methods=["POST"])
def take_screenshot():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["gowitness"]["check_cmd"]):
        return jsonify({"error": "gowitness n'est pas installé."}), 412

    # Assurez-vous que l'URL a un schéma (http:// ou https://) pour gowitness
    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "http://" + target_url

    # Generer un nom de fichier unique basé sur le hash de l'url
    import hashlib
    filename = hashlib.md5(target_url.encode()).hexdigest() + ".jpeg"
    filepath = os.path.join(SCREENSHOT_DIR, filename)

    cmd = ["gowitness", "scan", "single", "-u", target_url, "--screenshot-path", SCREENSHOT_DIR, "--screenshot-format", "jpeg"]

    try:
        # Lancement synchrone (prend quelques secondes)
        subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        
        # Gowitness enregistre le fichier sous forme de <protocole>-<domaine>.jpeg
        # Par defaut, on recupere l'image la plus recente dans le dossier si le nom est dur a deviner
        import glob
        list_of_files = glob.glob(os.path.join(SCREENSHOT_DIR, '*.jpeg'))
        if not list_of_files:
             return jsonify({"error": "La capture d'écran a échoué (fichier non généré)."}), 500
             
        latest_file = max(list_of_files, key=os.path.getctime)
        relative_path = os.path.relpath(latest_file, app.static_folder)
        # Fix backslashes on windows just in case
        relative_path = relative_path.replace('\\', '/')

        return jsonify({"screenshot_url": f"/static/{relative_path}"})

    except subprocess.TimeoutExpired:
        return jsonify({"error": "Délai dépassé pour la capture d'écran."}), 504
    except Exception as exc:
        return jsonify({"error": f"Erreur de capture : {str(exc)}"}), 500

# =====================================================================================
# SECTION 9 : WPSCAN
# =====================================================================================

WPSCAN_JOBS = {}

def _run_wpscan_thread(job_id, target_url):
    WPSCAN_JOBS[job_id]["status"] = "running"
    WPSCAN_JOBS[job_id]["log"] = []
    
    # URL needs scheme
    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "http://" + target_url

    cmd = ["wpscan", "--url", target_url, "--random-user-agent", "--format", "json", "--no-update"]
    cmd = _wrap_with_proxy(cmd)

    try:
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        
        # Wpscan exits with >0 if vulnerabilities are found or if an error occurs.
        # So we just try to parse stdout as JSON regardless of return code.
        results = {}
        try:
            # Look for JSON output. WPScan might output some text before JSON if there are warnings
            out = process.stdout.strip()
            # Find first '{'
            start_idx = out.find('{')
            if start_idx != -1:
                 results = json.loads(out[start_idx:])
        except json.JSONDecodeError as e:
            WPSCAN_JOBS[job_id]["status"] = "error"
            WPSCAN_JOBS[job_id]["error"] = f"Impossible de lire le JSON WPScan: {e}\nOut: {process.stdout[:100]}"
            return
            
        if not results:
             WPSCAN_JOBS[job_id]["status"] = "error"
             WPSCAN_JOBS[job_id]["error"] = "Aucun résultat exploitable retourné par WPScan."
             return

        WPSCAN_JOBS[job_id]["status"] = "done"
        WPSCAN_JOBS[job_id]["result"] = results
        save_history("wpscan", target_url, results)

    except subprocess.TimeoutExpired:
        WPSCAN_JOBS[job_id]["status"] = "error"
        WPSCAN_JOBS[job_id]["error"] = "Délai d'attente dépassé pour WPScan (5 min)."
    except Exception as exc:
        WPSCAN_JOBS[job_id]["status"] = "error"
        WPSCAN_JOBS[job_id]["error"] = f"Erreur inattendue : {exc}"


@app.route("/api/wpscan", methods=["POST"])
def start_wpscan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["wpscan"]["check_cmd"]):
        return jsonify({"error": "wpscan n'est pas installé."}), 412

    job_id = str(uuid.uuid4())
    WPSCAN_JOBS[job_id] = {
        "status": "queued",
        "result": None,
        "error": None,
        "log": [],
        "target_url": target_url,
        "started_at": datetime.now().isoformat(),
    }

    thread = threading.Thread(target=_run_wpscan_thread, args=(job_id, target_url), daemon=True)
    thread.start()

    return jsonify({"job_id": job_id})

@app.route("/api/wpscan-status/<job_id>", methods=["GET"])
def wpscan_status(job_id):
    job = WPSCAN_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)

# =====================================================================================
# SECTION 7.45 : AUDIT WEB — GOBUSTER, SQLMAP, JOOMSCAN, DROOPESCAN, MOODLESCAN
# =====================================================================================

AUDIT_JOBS = {}


def _run_gobuster_thread(job_id, target_url):
    """Lance Gobuster en arrière-plan pour la découverte de répertoires."""
    AUDIT_JOBS[job_id]["status"] = "running"
    AUDIT_JOBS[job_id]["output"] = []

    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "http://" + target_url

    # Wordlist standard Kali
    wordlist_candidates = [
        "/usr/share/wordlists/dirb/common.txt",
        "/usr/share/dirb/wordlists/common.txt",
        "/usr/share/wordlists/dirbuster/directory-list-2.3-small.txt",
    ]
    wordlist = next((w for w in wordlist_candidates if os.path.exists(w)), None)

    if not wordlist:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = "Aucune wordlist trouvée (dirb/dirbuster non installé)."
        return

    cmd = ["gobuster", "dir", "-u", target_url, "-w", wordlist, "-t", "20", "-q", "--no-progress", "-k"]
    cmd = _wrap_with_proxy(cmd)
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        lines = proc.stdout.splitlines()
        found = [l for l in lines if l.strip() and "(Status:" in l]
        AUDIT_JOBS[job_id]["status"] = "done"
        AUDIT_JOBS[job_id]["output"] = lines
        AUDIT_JOBS[job_id]["found_paths"] = found
        save_history("gobuster", target_url, {"paths": found[:100]})
    except subprocess.TimeoutExpired:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = "Délai dépassé pour Gobuster (5 min)."
    except Exception as e:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = str(e)


def _run_sqlmap_thread(job_id, target_url):
    """Lance SQLMap en mode batch sur l'URL cible, avec un pré-scan ParamSpider si possible."""
    AUDIT_JOBS[job_id]["status"] = "running"
    AUDIT_JOBS[job_id]["output"] = []

    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "http://" + target_url

    import urllib.parse
    domain = urllib.parse.urlparse(target_url).netloc
    if not domain:
        domain = target_url.split('/')[2] if '//' in target_url else target_url.split('/')[0]

    paramspider_output_file = f"/tmp/paramspider_{job_id}.txt"
    sqlmap_input_file = None

    # Etape 1: Tenter ParamSpider pour trouver des URLs avec paramètres
    AUDIT_JOBS[job_id]["output"].append(f"[*] Lancement de ParamSpider pour trouver des paramètres sur {domain}...")
    try:
        ps_cmd = ["paramspider", "-d", domain]
        subprocess.run(ps_cmd, capture_output=True, text=True, timeout=60, cwd="/tmp")
        # ParamSpider sauvegarde dans result/domain.txt ou results/domain.txt
        for results_dir in ["results", "result"]:
            expected_file = f"/tmp/{results_dir}/{domain}.txt"
            if os.path.exists(expected_file):
                import shutil
                shutil.copy(expected_file, paramspider_output_file)
                sqlmap_input_file = paramspider_output_file
                AUDIT_JOBS[job_id]["output"].append(f"[+] ParamSpider a trouvé des URLs. Transmission à SQLMap...")
                break
    except Exception as e:
        AUDIT_JOBS[job_id]["output"].append(f"[-] Erreur ParamSpider : {str(e)}. Fallback sur cible unique.")

    # Etape 2: Lancer SQLMap
    cmd = ["sqlmap", "--batch", "--level=1", "--risk=1",
           "--timeout=10", "--retries=1", "--output-dir=/tmp/sqlmap_out",
           "--disable-coloring", "--no-cast"]
    
    if sqlmap_input_file:
        cmd.extend(["-m", sqlmap_input_file])
    else:
        cmd.extend(["-u", target_url])

    cmd = _wrap_with_proxy(cmd)
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        lines = (proc.stdout + proc.stderr).splitlines()
        vulns = [l for l in lines if "injectable" in l.lower() or "Parameter:" in l or "[VULNERABLE]" in l or "[WARNING]" in l or "[INFO]" in l]
        AUDIT_JOBS[job_id]["status"] = "done"
        AUDIT_JOBS[job_id]["output"].extend(lines)
        AUDIT_JOBS[job_id]["vulnerabilities"] = [l for l in lines if "injectable" in l.lower() or "[VULNERABLE]" in l]
        AUDIT_JOBS[job_id]["summary_lines"] = vulns[:50]
        save_history("sqlmap", target_url, {"summary": vulns[:30]})
    except subprocess.TimeoutExpired:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = "Délai dépassé pour SQLMap (5 min)."
    except Exception as e:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = str(e)


def _run_cms_scanner_thread(job_id, cms_type, target_url):
    """Lance le scanner CMS approprié (wpscan, joomscan, droopescan, moodlescan)."""
    AUDIT_JOBS[job_id]["status"] = "running"
    AUDIT_JOBS[job_id]["output"] = []
    AUDIT_JOBS[job_id]["cms_type"] = cms_type

    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "http://" + target_url

    try:
        if cms_type == "wordpress":
            cmd = ["wpscan", "--url", target_url, "--random-user-agent", "--no-update",
                   "--enumerate", "p,t,u", "--format", "cli-no-colour"]
        elif cms_type == "joomla":
            cmd = ["joomscan", "--url", target_url]
        elif cms_type == "drupal":
            cmd = ["droopescan", "scan", "drupal", "-u", target_url]
        elif cms_type == "moodle":
            cmd = ["moodlescan", "-u", target_url]
        else:
            # Fallback générique droopescan
            cmd = ["droopescan", "scan", "-u", target_url]

        cmd = _wrap_with_proxy(cmd)
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        lines = (proc.stdout + proc.stderr).splitlines()
        AUDIT_JOBS[job_id]["status"] = "done"
        AUDIT_JOBS[job_id]["output"] = lines
        save_history(f"{cms_type}_scan", target_url, {"lines": lines[:200]})
    except subprocess.TimeoutExpired:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = f"Délai dépassé pour le scanner {cms_type} (5 min)."
    except Exception as e:
        AUDIT_JOBS[job_id]["status"] = "error"
        AUDIT_JOBS[job_id]["error"] = str(e)




@app.route("/api/audit/gobuster", methods=["POST"])
def start_gobuster():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()
    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400
    if not _tool_is_installed(["gobuster"]):
        return jsonify({"error": "gobuster n'est pas installé.", "not_installed": True}), 412
    job_id = str(uuid.uuid4())
    AUDIT_JOBS[job_id] = {"status": "queued", "output": [], "found_paths": [], "error": None, "target_url": target_url, "tool": "gobuster", "started_at": datetime.now().isoformat()}
    threading.Thread(target=_run_gobuster_thread, args=(job_id, target_url), daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/audit/sqlmap", methods=["POST"])
def start_sqlmap():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()
    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400
    if not _tool_is_installed(["sqlmap"]):
        return jsonify({"error": "sqlmap n'est pas installé.", "not_installed": True}), 412
    job_id = str(uuid.uuid4())
    AUDIT_JOBS[job_id] = {"status": "queued", "output": [], "vulnerabilities": [], "summary_lines": [], "error": None, "target_url": target_url, "tool": "sqlmap", "started_at": datetime.now().isoformat()}
    threading.Thread(target=_run_sqlmap_thread, args=(job_id, target_url), daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/audit/cms-scan", methods=["POST"])
def start_cms_scan():
    """Déclenche le scanner CMS approprié selon le CMS détecté par WhatWeb."""
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()
    cms_type = data.get("cms_type", "unknown").strip().lower()  # wordpress, joomla, drupal, moodle

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    # Vérification de la disponibilité du scanner approprié
    scanner_checks = {
        "wordpress": ["wpscan"],
        "joomla": ["joomscan"],
        "drupal": ["droopescan"],
        "moodle": ["moodlescan"],
    }
    check = scanner_checks.get(cms_type, ["droopescan"])
    if not _tool_is_installed(check):
        return jsonify({"error": f"Scanner {cms_type} ({check[0]}) n'est pas installé.", "not_installed": True}), 412

    job_id = str(uuid.uuid4())
    AUDIT_JOBS[job_id] = {"status": "queued", "output": [], "error": None, "target_url": target_url, "tool": f"{cms_type}_scan", "cms_type": cms_type, "started_at": datetime.now().isoformat()}
    threading.Thread(target=_run_cms_scanner_thread, args=(job_id, cms_type, target_url), daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/api/audit/status/<job_id>", methods=["GET"])
def audit_status(job_id):
    """Retourne le statut d'un job d'audit (Nikto/Gobuster/SQLMap/CMS)."""
    job = AUDIT_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)


# =====================================================================================
# SECTION 7.5 : WAFW00F (WAF DETECTION)
# =====================================================================================

@app.route("/api/waf-scan", methods=["POST"])
def waf_scan():
    """
    Exécute wafw00f pour détecter la présence d'un pare-feu applicatif.
    C'est un scan très rapide, donc on le fait de manière synchrone.
    """
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "URL cible manquante."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["wafw00f"]["check_cmd"]):
        return jsonify({"error": "wafw00f n'est pas installé."}), 412

    urls = []
    if not target_url.startswith("http"):
        urls = ["https://" + target_url]
    else:
        urls = [target_url]

    with tempfile.NamedTemporaryFile(suffix=".json", delete=False) as tmp_file:
        tmp_path = tmp_file.name

    cmd = ["wafw00f"] + urls + ["-f", "json", "-o", tmp_path]

    try:
        subprocess.run(cmd, capture_output=True, text=True, timeout=20)

        results = []
        if os.path.exists(tmp_path):
            with open(tmp_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
                if content:
                    results = json.loads(content)
            os.remove(tmp_path)

        # Structure du retour
        if not results:
            return jsonify({"detected": False, "firewall": None})

        # Wafw00f renvoie parfois une liste de dictionnaires si on a testé plusieurs URLs
        for waf_data in results:
            if waf_data.get("detected"):
                return jsonify({
                    "detected": True,
                    "firewall": waf_data.get("firewall"),
                    "manufacturer": waf_data.get("manufacturer")
                })
                
        return jsonify({"detected": False, "firewall": None})

    except subprocess.TimeoutExpired:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        return jsonify({"error": "Délai d'attente dépassé pour wafw00f."}), 504
    except Exception as exc:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        return jsonify({"error": f"Erreur lors de la détection WAF : {exc}"}), 500


# =====================================================================================
# SECTION 7.6 : WHOIS & DOMAIN RECON
# =====================================================================================

def extract_root_domain(target_url):
    cleaned = target_url.replace("http://", "").replace("https://", "").split("/")[0].split(":")[0].strip()
    parts = cleaned.split(".")
    if len(parts) >= 2:
        return ".".join(parts[-2:])
    return cleaned


@app.route("/api/whois-scan", methods=["POST"])
def whois_scan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "Cible manquante."}), 400

    domain = extract_root_domain(target_url)

    cmd = ["whois", domain]
    try:
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
        raw_text = process.stdout or process.stderr

        registrar = None
        created_date = None
        updated_date = None
        expiry_date = None
        dnssec = None
        statuses = []
        name_servers = []

        for line in raw_text.splitlines():
            line_str = line.strip()
            line_lower = line_str.lower()
            if not registrar and ("registrar:" in line_lower or "sponsoring registrar:" in line_lower):
                registrar = line_str.split(":", 1)[-1].strip()
            elif not created_date and ("creation date:" in line_lower or "created:" in line_lower or "created on:" in line_lower):
                created_date = line_str.split(":", 1)[-1].strip()
            elif not updated_date and ("updated date:" in line_lower or "last-update:" in line_lower or "changed:" in line_lower):
                updated_date = line_str.split(":", 1)[-1].strip()
            elif not expiry_date and ("registry expiry date:" in line_lower or "expiry date:" in line_lower or "expiration date:" in line_lower or "expires:" in line_lower):
                expiry_date = line_str.split(":", 1)[-1].strip()
            elif not dnssec and "dnssec:" in line_lower:
                dnssec = line_str.split(":", 1)[-1].strip()
            elif "domain status:" in line_lower or "status:" in line_lower:
                st = line_str.split(":", 1)[-1].strip().split()[0]
                if st and st not in statuses and len(statuses) < 4:
                    statuses.append(st)
            elif "name server:" in line_lower or "nserver:" in line_lower:
                ns = line_str.split(":", 1)[-1].strip().lower()
                if ns and ns not in name_servers:
                    name_servers.append(ns)

        return jsonify({
            "domain": domain,
            "registrar": registrar or "Non divulgué / Privé",
            "created_date": created_date or "Non renseignée",
            "updated_date": updated_date or "Non renseignée",
            "expiry_date": expiry_date or "Non renseignée",
            "dnssec": dnssec or "Inconnu / Non configuré",
            "statuses": statuses,
            "name_servers": name_servers[:6],
            "raw_text": raw_text[:4000]
        })
    except subprocess.TimeoutExpired:
        return jsonify({"error": "Délai d'attente dépassé pour la requête WHOIS."}), 504
    except Exception as exc:
        return jsonify({"error": f"Erreur WHOIS : {exc}"}), 500


# =====================================================================================
# SECTION 7.7 : HYBRID SUBDOMAIN ENUMERATION (SUBFINDER + SUBLIST3R + FAST DNS PROBE)
# =====================================================================================

@app.route("/api/subdomains-scan", methods=["POST"])
def subdomains_scan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "Cible manquante."}), 400

    root_domain = extract_root_domain(target_url)

    # 0. DÉTECTION DU WILDCARD DNS (Pour éliminer les faux positifs *.domaine)
    wildcard_ips = set()
    wildcard_detected = False
    for i in range(5):
        random_sub = f"wildcard-check-{uuid.uuid4().hex[:8]}.{root_domain}"
        try:
            w_ip = socket.gethostbyname(random_sub)
            if w_ip:
                wildcard_ips.add(w_ip)
                wildcard_detected = True
        except socket.gaierror:
            pass

    subdomains = set()

    # MOTEUR 1 : Subfinder (CT Logs, SecurityTrails, Chaos, DNSDumpster)
    try:
        sf_res = subprocess.run(["subfinder", "-d", root_domain, "-silent"], capture_output=True, text=True, timeout=25)
        for line in sf_res.stdout.splitlines():
            sd = line.strip().lower()
            if sd and root_domain in sd:
                subdomains.add(sd)
    except Exception:
        pass

    # MOTEUR 2 : Sublist3r OSINT (recherche passive multi-moteurs)
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp_file:
        tmp_path = tmp_file.name

    cmd = ["sublist3r", "-d", root_domain, "-o", tmp_path, "-t", "5"]
    try:
        subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        if os.path.exists(tmp_path):
            with open(tmp_path, 'r', encoding='utf-8', errors='ignore') as f:
                for line in f:
                    sd = line.strip().lower()
                    if sd and root_domain in sd:
                        subdomains.add(sd)
            os.remove(tmp_path)
    except Exception:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

    # MOTEUR 3 : theHarvester (OSINT étendu)
    # Déplacé dans une étape distincte (API /api/harvester-scan)

    # MOTEUR 4 : Sondage DNS direct multithreadé (avec filtrage anti-Wildcard)
    common_prefixes = [
        "www", "wiki", "mail", "webmail", "remote", "admin", "api", "dev", "test", "vpn",
        "cloud", "portal", "shop", "store", "git", "gitlab", "ns1", "ns2", "cpanel",
        "autodiscover", "blog", "crm", "matomo", "dolibarr", "app", "demo", "stage",
        "cyber", "leblog", "formation", "neoweb", "service-courtage", "ra"
    ]

    def resolve_prefix(prefix):
        fqdn = f"{prefix}.{root_domain}"
        try:
            ip = socket.gethostbyname(fqdn)
            if wildcard_detected and ip in wildcard_ips:
                return None
            return fqdn.lower()
        except socket.gaierror:
            return None

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=15) as executor:
            for res in executor.map(resolve_prefix, common_prefixes):
                if res:
                    subdomains.add(res)
    except Exception:
        pass

    # Filtrage final de tous les sous-domaines candidats contre les IP Wildcard
    filtered_subdomains = set()
    for sd in subdomains:
        try:
            ip = socket.gethostbyname(sd)
            if wildcard_detected and ip in wildcard_ips:
                continue
            filtered_subdomains.add(sd)
        except socket.gaierror:
            pass

    sorted_subdomains = sorted(list(filtered_subdomains))

    return jsonify({
        "domain": root_domain,
        "subdomains": sorted_subdomains,
        "count": len(sorted_subdomains),
        "wildcard_detected": wildcard_detected,
        "wildcard_ips": list(wildcard_ips)
    })


# =====================================================================================
# SECTION 7.7.5 : THEHARVESTER STANDALONE (LIVE STREAM)
# =====================================================================================

def _run_harvester_thread(job_id, target_url):
    SCAN_JOBS[job_id]["status"] = "running"
    SCAN_JOBS[job_id]["log"] = []
    
    root_domain = extract_root_domain(target_url)
    cmd = ["theHarvester", "-d", root_domain, "-b", "all", "-l", "100"]

    def append_log(line):
        log = SCAN_JOBS[job_id]["log"]
        log.append(line)
        if len(log) > MAX_LOG_LINES:
            del log[0:len(log) - MAX_LOG_LINES]

    try:
        run_cmd = cmd
        if shutil.which("stdbuf"):
            run_cmd = ["stdbuf", "-oL", "-eL"] + cmd

        process = subprocess.Popen(
            run_cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1
        )
        ACTIVE_PROCESSES[job_id] = process
    except Exception as e:
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = f"Erreur de lancement theHarvester: {e}"
        return

    def read_output():
        for line in process.stdout:
            clean_line = line.rstrip("\n")
            if clean_line:
                append_log(clean_line)

    t_out = threading.Thread(target=read_output, daemon=True)
    t_out.start()

    try:
        process.wait(timeout=900)
    except subprocess.TimeoutExpired:
        process.kill()
        process.wait()
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = "Délai d'attente dépassé pour theHarvester (15m)."
        append_log("[✘ Scan interrompu : délai dépassé]")
        return

    t_out.join(timeout=5)

    if process.returncode != 0:
        SCAN_JOBS[job_id]["status"] = "error"
        SCAN_JOBS[job_id]["error"] = "theHarvester a retourné une erreur."
        return

    subdomains = set()
    domain_pattern = re.compile(r'([a-zA-Z0-9.-]+\.' + re.escape(root_domain) + r')')
    for line in SCAN_JOBS[job_id]["log"]:
        for match in domain_pattern.finditer(line):
            sd = match.group(1).lower()
            if sd != root_domain:
                subdomains.add(sd)

    SCAN_JOBS[job_id]["status"] = "done"
    SCAN_JOBS[job_id]["result"] = {"subdomains": sorted(list(subdomains))}
    append_log(f"[✔ Scan theHarvester terminé, {len(subdomains)} sous-domaines trouvés]")


@app.route("/api/harvester-scan", methods=["POST"])
def harvester_scan():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"error": "Cible manquante."}), 400

    job_id = str(uuid.uuid4())
    SCAN_JOBS[job_id] = {"status": "pending", "log": [], "result": None}

    t = threading.Thread(target=_run_harvester_thread, args=(job_id, target_url), daemon=True)
    t.start()

    return jsonify({"job_id": job_id})


# =====================================================================================
# SECTION 7.8 : CHECK TARGET CONNECTIVITY (PING / HTTP CHECK FOR PAUSE MODAL)
# =====================================================================================

@app.route("/api/check-ping", methods=["POST"])
def check_ping():
    data = request.get_json(silent=True) or {}
    target_url = data.get("target_url", "").strip()

    if not target_url:
        return jsonify({"reachable": False, "error": "Cible manquante."})

    host = target_url.replace("http://", "").replace("https://", "").split("/")[0].split(":")[0]

    # Test 1 : Résolution DNS / Socket connect sur port 80 ou 443 ou ICMP ping
    try:
        ip = socket.gethostbyname(host)
    except socket.gaierror:
        return jsonify({"reachable": False, "error": f"Résolution DNS échouée pour {host}."})

    # Test 2 : Ping rapide 2 sec
    cmd = ["ping", "-c", "1", "-W", "2", host]
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode == 0:
        return jsonify({"reachable": True, "ip": ip, "method": "ICMP Ping"})

    # Test 3 : TCP connect port 80/443 si ICMP bloqué par pare-feu
    for port in [80, 443, 22, 8080]:
        try:
            with socket.create_connection((ip, port), timeout=2.5):
                return jsonify({"reachable": True, "ip": ip, "port": port, "method": f"TCP Port {port}"})
        except OSError:
            pass

    return jsonify({"reachable": False, "ip": ip, "error": "Le serveur ne répond ni au Ping ICMP ni aux connexions TCP (80/443)."})



# =====================================================================================
# SECTION 7.8 : NETDISCOVER (ARP DISCOVERY)
# =====================================================================================

NETDISCOVER_JOBS = {}

CONFIGURED_SUDO_PASSWORD = ""

@app.route("/api/settings/sudo", methods=["POST"])
def set_sudo_password():
    global CONFIGURED_SUDO_PASSWORD
    data = request.get_json(silent=True) or {}
    password = data.get("password", "")

    if not password:
        CONFIGURED_SUDO_PASSWORD = ""
        return jsonify({"success": True, "message": "Mot de passe effacé."})

    try:
        res = subprocess.run(
            ["sudo", "-k", "-S", "id"],
            input=f"{password}\n",
            capture_output=True,
            text=True,
            timeout=5
        )
        if res.returncode == 0:
            CONFIGURED_SUDO_PASSWORD = password
            return jsonify({"success": True, "message": "Accès sudo vérifié et activé pour Netdiscover !" })
        else:
            return jsonify({"success": False, "error": "Mot de passe sudo/root invalide."}), 401
    except Exception as exc:
        return jsonify({"success": False, "error": f"Erreur de vérification sudo : {exc}"}), 500


@app.route("/api/settings/sudo/clear", methods=["POST"])
def clear_sudo_password():
    global CONFIGURED_SUDO_PASSWORD
    CONFIGURED_SUDO_PASSWORD = ""
    return jsonify({"success": True, "message": "Mot de passe effacé."})


def _run_netdiscover_thread(job_id, target):
    NETDISCOVER_JOBS[job_id]["status"] = "running"
    NETDISCOVER_JOBS[job_id]["log"] = []

    # Clean target for Netdiscover
    is_valid, msg = validate_target(target)
    if is_valid and msg:
        target = msg[0]


    if CONFIGURED_SUDO_PASSWORD:
        cmd = ["sudo", "-S", "netdiscover", "-r", target, "-P", "-N", "-f"]
        input_data = f"{CONFIGURED_SUDO_PASSWORD}\n"
    else:
        cmd = ["netdiscover", "-r", target, "-P", "-N", "-f"]
        input_data = None

    try:
        process = subprocess.run(cmd, input=input_data, capture_output=True, text=True, timeout=40)
        output = (process.stdout or "") + (process.stderr or "")

        if "permission" in output.lower() or "cap_net_raw" in output.lower() or "socket failed" in output.lower():
            # Fallback vers ARP scan via nmap -sn -PR
            fallback_cmd = ["nmap", "-sn", "-PR", target, "-oX", "-"]
            fb_process = subprocess.run(fallback_cmd, capture_output=True, text=True, timeout=30)
            if fb_process.returncode == 0 and fb_process.stdout:
                parsed_hosts = parse_nmap_xml(fb_process.stdout)
                devices = []
                for h in parsed_hosts:
                    devices.append({
                        "ip": h.get("ip"),
                        "mac": h.get("mac") or "Inconnu",
                        "count": 1,
                        "len": 60,
                        "vendor": h.get("vendor") or "Inconnu"
                    })
                NETDISCOVER_JOBS[job_id]["status"] = "done"
                NETDISCOVER_JOBS[job_id]["result"] = devices
                save_history("netdiscover", target, devices)
                return

            NETDISCOVER_JOBS[job_id]["status"] = "error"
            NETDISCOVER_JOBS[job_id]["error"] = (
                "Permission refusée par le système Linux (capture de paquets ARP bruts). "
                "Pour autoriser Netdiscover dans l'application Flask, exécutez la commande suivante dans votre terminal : "
                "sudo setcap cap_net_raw,cap_net_admin=eip /usr/sbin/netdiscover"
            )
            return

        devices = []
        # Netdiscover format: IP MAC Count Len Vendor
        pattern = re.compile(r"^\s*([\d\.]+)\s+([0-9a-fA-F:]{17})\s+(\d+)\s+(\d+)\s+(.*)$")
        for line in output.splitlines():
            line_str = line.strip()
            match = pattern.match(line_str)
            if match:
                ip_val = match.group(1)
                mac_val = match.group(2)
                v_val = match.group(5).strip() or "Inconnu"
                identity = resolve_host_identity(ip_val, mac_val, v_val)

                devices.append({
                    "ip": ip_val,
                    "mac": mac_val,
                    "count": int(match.group(3)),
                    "len": int(match.group(4)),
                    "vendor": identity["vendor"],
                    "hostname": identity["hostname"] or identity["netbios_name"],
                    "netbios_name": identity["netbios_name"],
                    "workgroup": identity["workgroup"],
                    "category": identity["category"],
                    "icon": identity["icon"]
                })

        NETDISCOVER_JOBS[job_id]["status"] = "done"
        NETDISCOVER_JOBS[job_id]["result"] = devices
        save_history("netdiscover", target, devices)

    except subprocess.TimeoutExpired:
        NETDISCOVER_JOBS[job_id]["status"] = "error"
        NETDISCOVER_JOBS[job_id]["error"] = "Délai d'attente dépassé pour Netdiscover (40s)."
    except Exception as exc:
        NETDISCOVER_JOBS[job_id]["status"] = "error"
        NETDISCOVER_JOBS[job_id]["error"] = f"Erreur lors de l'exécution de Netdiscover : {exc}"


@app.route("/api/netdiscover-scan", methods=["POST"])
def start_netdiscover_scan():
    data = request.get_json(silent=True) or {}
    target = data.get("target", "").strip()

    if not target:
        return jsonify({"error": "Cible manquante."}), 400

    if not _tool_is_installed(REQUIRED_TOOLS["netdiscover"]["check_cmd"]):
        return jsonify({"error": "netdiscover n'est pas installé."}), 412

    job_id = str(uuid.uuid4())
    NETDISCOVER_JOBS[job_id] = {
        "status": "queued",
        "result": None,
        "error": None,
        "target": target,
        "started_at": datetime.now().isoformat(),
    }

    thread = threading.Thread(target=_run_netdiscover_thread, args=(job_id, target), daemon=True)
    thread.start()

    return jsonify({"job_id": job_id})


@app.route("/api/netdiscover-status/<job_id>", methods=["GET"])
def netdiscover_status(job_id):
    job = NETDISCOVER_JOBS.get(job_id)
    if job is None:
        return jsonify({"error": "job_id inconnu"}), 404
    return jsonify(job)


# =====================================================================================
# SECTION 8 : ROUTES FRONTEND
# =====================================================================================

@app.route("/")
def index():
    return render_template("home.html", active_page="home")

@app.route("/web")
def web():
    return render_template("web.html", active_page="web")

@app.route("/local")
def local():
    return render_template("local.html", active_page="local")

@app.route("/history")
@app.route("/history_page")
def history_page():
    return render_template("history.html", active_page="history")

@app.route("/settings")
def settings():
    return render_template("settings.html", active_page="settings")

@app.route("/api/history", methods=["GET"])
def get_history():
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT id, tool, target, mode, timestamp FROM history ORDER BY timestamp DESC").fetchall()
            return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/<int:history_id>", methods=["GET"])
def get_history_detail(history_id):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM history WHERE id = ?", (history_id,)).fetchone()
            if row:
                data = dict(row)
                data["result_json"] = json.loads(data["result_json"])
                return jsonify(data)
            return jsonify({"error": "Non trouvé"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/<int:history_id>", methods=["DELETE"])
def delete_history(history_id):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("DELETE FROM history WHERE id = ?", (history_id,))
            conn.commit()
            return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/history/delete-batch", methods=["POST"])
def delete_history_batch():
    try:
        data = request.get_json() or {}
        ids = data.get("ids", [])
        if not ids:
            return jsonify({"error": "Aucun rapport sélectionné"}), 400
        
        with sqlite3.connect(DB_PATH) as conn:
            placeholders = ",".join(["?"] * len(ids))
            conn.execute(f"DELETE FROM history WHERE id IN ({placeholders})", ids)
            conn.commit()
            return jsonify({"success": True, "deleted_count": len(ids)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# =====================================================================================
# SECTION 9 : GESTION DES RAPPORTS PERSONNALISÉS (API & EXPORTATION)
# =====================================================================================

@app.route("/reports")
def reports_page():
    return render_template("reports.html", active_page="reports")

@app.route("/api/reports", methods=["GET"])
def get_reports():
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT id, title, author, description, created_at, updated_at FROM custom_reports ORDER BY updated_at DESC").fetchall()
            return jsonify([dict(r) for r in rows])
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports", methods=["POST"])
def create_report():
    try:
        data = request.get_json() or {}
        title = data.get("title", "Nouveau Rapport").strip() or "Nouveau Rapport"
        author = data.get("author", "Auditeur Cybersécurité").strip()
        description = data.get("description", "").strip()
        initial_html = f"<h1>{title}</h1><p><strong>Auditeur :</strong> {author}</p><p><strong>Date :</strong> {datetime.now().strftime('%d/%m/%Y %H:%M')}</p><hr><p>{description}</p>"
        
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO custom_reports (title, author, description, content_html, updated_at) VALUES (?, ?, ?, ?, CURRENT_TIMESTAMP)",
                (title, author, description, initial_html)
            )
            report_id = cursor.lastrowid
            conn.commit()
            return jsonify({"id": report_id, "title": title, "success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports/<int:report_id>", methods=["GET"])
def get_report_detail(report_id):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM custom_reports WHERE id = ?", (report_id,)).fetchone()
            if not row:
                return jsonify({"error": "Rapport non trouvé"}), 404
            items = conn.execute("SELECT * FROM report_items WHERE report_id = ? ORDER BY added_at ASC", (report_id,)).fetchall()
            data = dict(row)
            data["items"] = [dict(i) for i in items]
            return jsonify(data)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports/<int:report_id>", methods=["PUT"])
def update_report(report_id):
    try:
        data = request.get_json() or {}
        title = data.get("title")
        author = data.get("author")
        description = data.get("description")
        content_html = data.get("content_html")

        fields = []
        params = []
        if title is not None:
            fields.append("title = ?")
            params.append(title)
        if author is not None:
            fields.append("author = ?")
            params.append(author)
        if description is not None:
            fields.append("description = ?")
            params.append(description)
        if content_html is not None:
            fields.append("content_html = ?")
            params.append(content_html)

        if not fields:
            return jsonify({"error": "Aucune modification fournie"}), 400

        fields.append("updated_at = CURRENT_TIMESTAMP")
        params.append(report_id)

        with sqlite3.connect(DB_PATH) as conn:
            conn.execute(f"UPDATE custom_reports SET {', '.join(fields)} WHERE id = ?", params)
            conn.commit()
            return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports/<int:report_id>", methods=["DELETE"])
def delete_report(report_id):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute("DELETE FROM custom_reports WHERE id = ?", (report_id,))
            conn.execute("DELETE FROM report_items WHERE report_id = ?", (report_id,))
            conn.commit()
            return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports/<int:report_id>/append", methods=["POST"])
def append_to_report(report_id):
    try:
        data = request.get_json() or {}
        item_title = data.get("title", "Élément d'analyse")
        item_html = data.get("html", "")

        if not item_html:
            return jsonify({"error": "Contenu HTML vide"}), 400

        with sqlite3.connect(DB_PATH) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT content_html FROM custom_reports WHERE id = ?", (report_id,)).fetchone()
            if not row:
                return jsonify({"error": "Rapport non trouvé"}), 404

            current_html = row["content_html"] or ""
            new_block = f'<div class="report-section-block" style="margin-top:20px; padding:15px; border-left:4px solid #00ff41; background:rgba(0,255,65,0.05);"><h3>{item_title}</h3>{item_html}</div>'
            updated_html = current_html + new_block

            conn.execute(
                "UPDATE custom_reports SET content_html = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (updated_html, report_id)
            )
            conn.execute(
                "INSERT INTO report_items (report_id, item_title, item_html) VALUES (?, ?, ?)",
                (report_id, item_title, item_html)
            )
            conn.commit()
            return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/reports/<int:report_id>/export/docx", methods=["GET"])
def export_report_docx(report_id):
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from bs4 import BeautifulSoup
    except ImportError:
        return jsonify({"error": "Module python-docx non disponible"}), 500

    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        report = conn.execute("SELECT * FROM custom_reports WHERE id = ?", (report_id,)).fetchone()
        if not report:
            return jsonify({"error": "Rapport introuvable"}), 404

    doc = docx.Document()
    title_p = doc.add_heading(report["title"], level=0)
    meta_p = doc.add_paragraph(f"Auditeur : {report['author']} | Date : {report['updated_at']}")
    meta_p.runs[0].font.italic = True
    meta_p.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    html_content = report["content_html"] or ""
    temp_files = []
    
    try:
        soup = BeautifulSoup(html_content, "html.parser")
        for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'img']):
            if elem.name in ['h1', 'h2', 'h3', 'h4']:
                level = int(elem.name[1])
                doc.add_heading(elem.get_text().strip(), level=level)
            elif elem.name == 'p':
                text = elem.get_text().strip()
                if text:
                    doc.add_paragraph(text)
            elif elem.name == 'table':
                rows = elem.find_all('tr')
                if rows:
                    col_count = max(len(r.find_all(['td', 'th'])) for r in rows)
                    if col_count > 0:
                        doc_table = doc.add_table(rows=len(rows), cols=col_count)
                        doc_table.style = 'Table Grid'
                        for r_idx, r in enumerate(rows):
                            cols = r.find_all(['td', 'th'])
                            for c_idx, c in enumerate(cols):
                                if c_idx < col_count:
                                    doc_table.cell(r_idx, c_idx).text = c.get_text().strip()
            elif elem.name == 'img' and elem.get('src', '').startswith('data:image'):
                src = elem['src']
                try:
                    header, encoded = src.split(',', 1)
                    img_bytes = base64.b64decode(encoded)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                        tmp_img.write(img_bytes)
                        tmp_path = tmp_img.name
                        temp_files.append(tmp_path)
                    doc.add_picture(tmp_path, width=Inches(6.0))
                except Exception as img_err:
                    print(f"Erreur insertion image Word : {img_err}")
    except Exception as parse_err:
        for line in html_content.replace("<br>", "\n").replace("</p>", "\n").split("\n"):
            clean_text = re.sub('<[^<]+?>', '', line).strip()
            if clean_text:
                doc.add_paragraph(clean_text)

    mem_file = io.BytesIO()
    doc.save(mem_file)

    for tmp in temp_files:
        if os.path.exists(tmp):
            try:
                os.remove(tmp)
            except Exception:
                pass

    mem_file.seek(0)
    filename = f"Rapport_{report_id}_{re.sub(r'[^a-zA-Z0-9]', '_', report['title'])}.docx"
    return send_file(
        mem_file,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )

@app.route("/api/reports/<int:report_id>/export/pdf", methods=["GET"])
def export_report_pdf(report_id):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from bs4 import BeautifulSoup
    except ImportError:
        return jsonify({"error": "Module reportlab non disponible"}), 500

    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        report = conn.execute("SELECT * FROM custom_reports WHERE id = ?", (report_id,)).fetchone()
        if not report:
            return jsonify({"error": "Rapport introuvable"}), 404

    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=18, textColor=colors.HexColor('#00aa2b'), spaceAfter=10)
    meta_style = ParagraphStyle('MetaStyle', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#555555'), spaceAfter=15)
    body_style = ParagraphStyle('BodyStyle', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=8)

    story = []
    story.append(Paragraph(report['title'], title_style))
    story.append(Paragraph(f"Auditeur : {report['author']} | Date : {report['updated_at']}", meta_style))
    story.append(Spacer(1, 10))

    html_content = report["content_html"] or ""
    temp_files = []

    try:
        soup = BeautifulSoup(html_content, "html.parser")
        for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'table', 'img']):
            if elem.name in ['h1', 'h2', 'h3', 'h4']:
                level_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=13, textColor=colors.HexColor('#1a5e20'), spaceBefore=8, spaceAfter=4)
                story.append(Paragraph(elem.get_text().strip(), level_style))
            elif elem.name == 'p':
                text = elem.get_text().strip()
                if text:
                    story.append(Paragraph(text, body_style))
            elif elem.name == 'table':
                rows = elem.find_all('tr')
                table_data = []
                for r in rows:
                    cols = r.find_all(['td', 'th'])
                    row_data = [Paragraph(c.get_text().strip(), body_style) for c in cols]
                    if row_data:
                        table_data.append(row_data)
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#e8f5e9')),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#1b5e20')),
                        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cccccc')),
                        ('PADDING', (0,0), (-1,-1), 4),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 10))
            elif elem.name == 'img' and elem.get('src', '').startswith('data:image'):
                src = elem['src']
                try:
                    header, encoded = src.split(',', 1)
                    img_bytes = base64.b64decode(encoded)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                        tmp_img.write(img_bytes)
                        tmp_path = tmp_img.name
                        temp_files.append(tmp_path)
                    rl_img = RLImage(tmp_path, width=450, height=250)
                    story.append(rl_img)
                    story.append(Spacer(1, 10))
                except Exception as img_err:
                    print(f"Erreur insertion image PDF : {img_err}")
    except Exception:
        clean_text = re.sub('<[^<]+?>', '', html_content)
        story.append(Paragraph(clean_text, body_style))

    doc.build(story)

    for tmp in temp_files:
        if os.path.exists(tmp):
            try:
                os.remove(tmp)
            except Exception:
                pass

    pdf_buffer.seek(0)
    filename = f"Rapport_{report_id}_{re.sub(r'[^a-zA-Z0-9]', '_', report['title'])}.pdf"
    return send_file(
        pdf_buffer,
        mimetype="application/pdf",
        as_attachment=True,
        download_name=filename
    )


def is_port_free(port, host="127.0.0.1"):
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        try:
            s.bind((host, port))
            return True
        except OSError:
            return False


def find_available_port(preferred_port=5000, start_fallback=10001, host="127.0.0.1"):
    env_port = os.environ.get("PORT")
    if env_port:
        try:
            p = int(env_port)
            if is_port_free(p, host):
                return p
        except ValueError:
            pass

    if is_port_free(preferred_port, host):
        return preferred_port

    print(f"[!] Le port {preferred_port} est occupé. Recherche d'un port disponible >= {start_fallback}...")
    p = start_fallback
    while p < 65535:
        if is_port_free(p, host):
            print(f"[+] Port disponible sélectionné : {p}")
            return p
        p += 1

    return preferred_port


# =====================================================================================
# SECTION ONIONHOP — Démarrage / Arrêt / Statut depuis le dashboard
# =====================================================================================

_ONIONHOP_PROCESS = None  # Référence globale au processus OnionHop en cours

def _wrap_with_proxy(cmd, use_sudo=False):
    """
    Si OnionHop (Tor) est actif, encapsule la commande avec proxychains4.
    Prend en compte sudo car 'proxychains sudo' échoue à cause du nettoyage LD_PRELOAD.
    Il faut faire 'sudo proxychains'.
    """
    global _ONIONHOP_PROCESS
    if _ONIONHOP_PROCESS and _ONIONHOP_PROCESS.poll() is None:
        import shutil
        pc = "proxychains4" if shutil.which("proxychains4") else ("proxychains" if shutil.which("proxychains") else None)
        if pc:
            if use_sudo:
                # ex: on a ["sudo", "-S", "nmap"] -> ["sudo", "-S", "proxychains4", "-q", "nmap"]
                return ["sudo", "-S", pc, "-q"] + cmd[2:]
            else:
                return [pc, "-q"] + cmd
    return cmd

def _find_onionhop_binary():
    """Retourne le chemin du binaire OnionHop si disponible, sinon None."""
    home_dir = os.path.expanduser("~")
    candidates = [
        os.path.join(home_dir, ".local", "bin", "OnionHop-x86_64.AppImage"),
        os.path.join(home_dir, ".local", "bin", "onionhop"),
        "/usr/local/bin/OnionHop-x86_64.AppImage",
        "onionhop",
        "OnionHop",
    ]
    for c in candidates:
        path = shutil.which(c) if not os.path.isabs(c) else (c if os.access(c, os.X_OK) else None)
        if path:
            return path
    return None


@app.route("/api/onionhop/start", methods=["POST"])
def onionhop_start():
    global _ONIONHOP_PROCESS
    # Vérifier si déjà lancé
    if _ONIONHOP_PROCESS and _ONIONHOP_PROCESS.poll() is None:
        return jsonify({"status": "already_running", "message": "OnionHop est déjà en cours d'exécution."})
    path = _find_onionhop_binary()
    if not path:
        return jsonify({"status": "not_found", "message": "OnionHop introuvable sur ce système."}), 404
    try:
        env = os.environ.copy()
        env["DISPLAY"] = env.get("DISPLAY", ":0")
        _ONIONHOP_PROCESS = subprocess.Popen([path], env=env, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return jsonify({"status": "started", "pid": _ONIONHOP_PROCESS.pid, "message": f"OnionHop lancé (PID {_ONIONHOP_PROCESS.pid})."})
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route("/api/onionhop/stop", methods=["POST"])
def onionhop_stop():
    global _ONIONHOP_PROCESS
    # Tuer aussi les processus OnionHop restants
    try:
        subprocess.run(["pkill", "-f", "OnionHop"], capture_output=True)
    except Exception:
        pass
    if _ONIONHOP_PROCESS:
        try:
            _ONIONHOP_PROCESS.terminate()
            _ONIONHOP_PROCESS.wait(timeout=3)
        except Exception:
            try:
                _ONIONHOP_PROCESS.kill()
            except Exception:
                pass
        _ONIONHOP_PROCESS = None
    return jsonify({"status": "stopped", "message": "OnionHop arrêté."})


@app.route("/api/onionhop/status", methods=["GET"])
def onionhop_status():
    global _ONIONHOP_PROCESS
    running = _ONIONHOP_PROCESS is not None and _ONIONHOP_PROCESS.poll() is None
    # Récupérer l'IP publique actuelle (via Tor si actif, sinon IP réelle)
    tor_ip = None
    if running:
        try:
            import socks
            s = socks.socksocket()
            s.set_proxy(socks.SOCKS5, "127.0.0.1", 9050)
            s.settimeout(5)
            s.connect(("api.ipify.org", 80))
            s.send(b"GET /?format=json HTTP/1.0\r\nHost: api.ipify.org\r\n\r\n")
            resp = s.recv(4096).decode("utf-8", errors="ignore")
            s.close()
            import re
            m = re.search(r'"ip"\s*:\s*"([\d\.]+)"', resp)
            if m:
                tor_ip = m.group(1)
        except Exception:
            tor_ip = None
    return jsonify({"running": running, "tor_ip": tor_ip})


def open_onionhop_or_default_browser(url="http://127.0.0.1:5000"):
    """
    Lance OnionHop 3.7.8 (si présent) et ouvre toujours Matrix Scanner dans le navigateur par défaut.
    """
    time.sleep(1.2)
    # Lancement prioritaire d'OnionHop si disponible
    home_dir = os.path.expanduser("~")
    onion_candidates = [
        "onionhop",
        "onionhop-3.7.8",
        "onion-hop",
        "OnionHop",
        "onionhop3.7.8",
        os.path.join(home_dir, ".local", "bin", "OnionHop-x86_64.AppImage"),
        os.path.join(home_dir, ".local", "bin", "onionhop"),
        "/usr/local/bin/OnionHop-x86_64.AppImage",
    ]
    for cmd in onion_candidates:
        path = shutil.which(cmd) if not os.path.isabs(cmd) else (cmd if os.access(cmd, os.X_OK) else None)
        if path:
            try:
                subprocess.Popen([path])
                print(f"[+] Application OnionHop lancée ({path})")
                break
            except Exception as e:
                print(f"[!] Échec du lancement d'OnionHop avec {cmd} : {e}")

    # Ouverture de Matrix Scanner dans le navigateur par défaut
    try:
        import webbrowser
        webbrowser.open(url)
        print(f"[+] Matrix Scanner ouvert dans le navigateur par défaut ({url})")
    except Exception as e:
        print(f"[!] Erreur d'ouverture du navigateur par défaut : {e}")


if __name__ == "__main__":
    server_port = find_available_port(5000, 10001)
    app_url = f"http://127.0.0.1:{server_port}"

    # Lancement d'OnionHop 3.7.8 / navigateur par défaut en arrière-plan
    threading.Thread(target=open_onionhop_or_default_browser, args=(app_url,), daemon=True).start()

    # debug=False — empêche le fork du processus fils Werkzeug (stable en arrière-plan)
    # use_reloader=False — empêche les redémarrages lors des écritures SQLite
    print(f"[+] Démarrage de Matrix Scanner sur {app_url}")
    # Expose to Docker container network if FLASK_HOST is provided, otherwise bind to localhost
    flask_host = os.environ.get("FLASK_HOST", "127.0.0.1")
    app.run(host=flask_host, port=server_port, debug=False, use_reloader=False)
